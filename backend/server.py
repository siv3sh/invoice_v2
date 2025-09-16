from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import StreamingResponse, JSONResponse, Response
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, validator
from typing import List, Optional, Dict, Any
import os
import logging
import uuid
import io
import pandas as pd
import openpyxl
from datetime import datetime, timedelta
from decimal import Decimal
import re
from pathlib import Path
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from io import BytesIO
from enum import Enum
import bcrypt
import jwt
from fastapi.middleware.cors import CORSMiddleware
import pdfplumber
import PyPDF2
import tabula
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.layout import LAParams
import docx

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# FastAPI app
app = FastAPI(title="Activus Invoice Management System", version="1.0.0")
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()
SECRET_KEY = "activus_secret_key_2024"

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# PDF Text Extraction Models
class POExtractedData(BaseModel):
    po_number: Optional[str] = None
    po_date: Optional[str] = None
    vendor_name: Optional[str] = None
    client_name: Optional[str] = None
    total_amount: Optional[float] = None
    line_items: List[Dict[str, Any]] = []
    terms_conditions: Optional[str] = None
    delivery_date: Optional[str] = None
    contact_info: Optional[Dict[str, str]] = {}
    raw_text: Optional[str] = None
    extraction_method: Optional[str] = None
    confidence_score: Optional[float] = None

class POPDFParser:
    """Comprehensive PDF text extraction engine for Purchase Orders"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
        self.extraction_methods = ['pdfplumber', 'pdfminer', 'pypdf2', 'tabula']
    
    async def extract_from_file(self, file_content: bytes, original_filename: str) -> POExtractedData:
        """Main method to extract data from uploaded file"""
        file_extension = Path(original_filename).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_content, original_filename)
            elif file_extension == '.docx':
                return await self._extract_from_docx(file_content, original_filename)
        except Exception as e:
            logger.error(f"Error extracting from {original_filename}: {str(e)}")
            return POExtractedData(
                raw_text=f"Extraction failed: {str(e)}",
                extraction_method="error",
                confidence_score=0.0
            )
    
    async def _extract_from_pdf(self, file_content: bytes, filename: str) -> POExtractedData:
        """Extract data from PDF using multiple methods"""
        pdf_file = BytesIO(file_content)
        results = []
        
        # Method 1: PDFPlumber (best for structured PDFs)
        try:
            with pdfplumber.open(pdf_file) as pdf:
                text_content = ""
                tables = []
                for page in pdf.pages:
                    text_content += page.extract_text() or ""
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                
                result = self._parse_po_text(text_content, "pdfplumber")
                result.line_items = self._extract_line_items_from_tables(tables)
                results.append(result)
        except Exception as e:
            logger.warning(f"PDFPlumber failed for {filename}: {str(e)}")
        
        # Method 2: PDFMiner (good for complex layouts)
        try:
            pdf_file.seek(0)
            text_content = pdfminer_extract_text(pdf_file, laparams=LAParams())
            result = self._parse_po_text(text_content, "pdfminer")
            results.append(result)
        except Exception as e:
            logger.warning(f"PDFMiner failed for {filename}: {str(e)}")
        
        # Method 3: PyPDF2 (fallback)
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text()
            result = self._parse_po_text(text_content, "pypdf2")
            results.append(result)
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {filename}: {str(e)}")
        
        # Method 4: Tabula (for table-heavy PDFs)
        try:
            pdf_file.seek(0)
            # Save to temporary file for tabula
            temp_file_path = f"/tmp/{filename}"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(file_content)
            
            tables = tabula.read_pdf(temp_file_path, pages='all', multiple_tables=True)
            if tables:
                result = POExtractedData(
                    line_items = self._extract_line_items_from_dataframes(tables),
                    extraction_method="tabula",
                    confidence_score=0.7
                )
                results.append(result)
            
            # Clean up temp file
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"Tabula failed for {filename}: {str(e)}")
        
        # Return the best result
        return self._select_best_result(results)
    
    async def _extract_from_docx(self, file_content: bytes, filename: str) -> POExtractedData:
        """Extract data from DOCX file"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            result = self._parse_po_text(text_content, "docx")
            result.line_items = self._extract_line_items_from_tables(tables_data)
            return result
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX {filename}: {str(e)}")
            return POExtractedData(
                raw_text=f"DOCX extraction failed: {str(e)}",
                extraction_method="docx_error",
                confidence_score=0.0
            )
    
    def _parse_po_text(self, text: str, method: str) -> POExtractedData:
        """Parse extracted text to identify PO components"""
        result = POExtractedData(raw_text=text, extraction_method=method)
        
        # PO Number patterns
        po_patterns = [
            r'PO\s*(?:NUMBER|NO|#)?\s*:?\s*([A-Z0-9\-\/]+)',
            r'PURCHASE\s*ORDER\s*(?:NUMBER|NO|#)?\s*:?\s*([A-Z0-9\-\/]+)',
            r'ORDER\s*(?:NUMBER|NO|#)?\s*:?\s*([A-Z0-9\-\/]+)'
        ]
        
        for pattern in po_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result.po_number = match.group(1).strip()
                break
        
        # Date patterns
        date_patterns = [
            r'PO\s*DATE\s*:?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'ORDER\s*DATE\s*:?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'DATE\s*:?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result.po_date = match.group(1).strip()
                break
        
        # Vendor/Supplier patterns
        vendor_patterns = [
            r'VENDOR\s*(?:NAME)?\s*:?\s*([^\n]+)',
            r'SUPPLIER\s*(?:NAME)?\s*:?\s*([^\n]+)',
            r'TO\s*:?\s*([^\n]+)'
        ]
        
        for pattern in vendor_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                vendor_line = match.group(1).strip()
                if len(vendor_line) > 3 and not re.match(r'^\d+', vendor_line):
                    result.vendor_name = vendor_line[:100]  # Limit length
                    break
        
        # Client/Company patterns
        client_patterns = [
            r'FROM\s*:?\s*([^\n]+)',
            r'COMPANY\s*(?:NAME)?\s*:?\s*([^\n]+)',
            r'CLIENT\s*(?:NAME)?\s*:?\s*([^\n]+)'
        ]
        
        for pattern in client_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                client_line = match.group(1).strip()
                if len(client_line) > 3:
                    result.client_name = client_line[:100]
                    break
        
        # Total amount patterns
        amount_patterns = [
            r'TOTAL\s*(?:AMOUNT)?\s*:?\s*[₹\$]?\s*([\d,]+\.?\d*)',
            r'GRAND\s*TOTAL\s*:?\s*[₹\$]?\s*([\d,]+\.?\d*)',
            r'NET\s*AMOUNT\s*:?\s*[₹\$]?\s*([\d,]+\.?\d*)'
        ]
        
        for pattern in amount_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                amount_str = match.group(1).replace(',', '')
                try:
                    result.total_amount = float(amount_str)
                    break
                except ValueError:
                    continue
        
        # Delivery date patterns
        delivery_patterns = [
            r'DELIVERY\s*(?:DATE|BY)?\s*:?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'EXPECTED\s*DELIVERY\s*:?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})'
        ]
        
        for pattern in delivery_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result.delivery_date = match.group(1).strip()
                break
        
        # Contact info patterns
        email_pattern = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        phone_pattern = r'(\+?\d{1,3}[\s\-]?\(?\d{3,}\)?[\s\-]?\d{3,}[\s\-]?\d{3,})'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        
        result.contact_info = {}
        if emails:
            result.contact_info['email'] = emails[0]
        if phones:
            result.contact_info['phone'] = phones[0]
        
        # Calculate confidence score
        score = 0
        if result.po_number: score += 25
        if result.vendor_name: score += 20
        if result.total_amount: score += 20
        if result.po_date: score += 15
        if result.line_items: score += 20
        
        result.confidence_score = min(score / 100.0, 1.0)
        
        return result
    
    def _extract_line_items_from_tables(self, tables: List[List[List[str]]]) -> List[Dict[str, Any]]:
        """Extract line items from table data"""
        line_items = []
        
        for table in tables:
            if not table or len(table) < 2:
                continue
            
            # Assume first row is header
            headers = [h.lower().strip() for h in table[0]]
            
            # Map common header variations
            header_mapping = {
                'description': ['description', 'item', 'product', 'service'],
                'quantity': ['quantity', 'qty', 'nos', 'number'],
                'rate': ['rate', 'price', 'unit price', 'cost'],
                'amount': ['amount', 'total', 'value', 'line total'],
                'unit': ['unit', 'uom', 'u.o.m', 'measure']
            }
            
            # Find column indices
            col_indices = {}
            for key, variations in header_mapping.items():
                for i, header in enumerate(headers):
                    if any(var in header for var in variations):
                        col_indices[key] = i
                        break
            
            # Extract data rows
            for row in table[1:]:
                if len(row) < 2:
                    continue
                
                item = {}
                
                # Extract based on found columns
                if 'description' in col_indices and col_indices['description'] < len(row):
                    item['description'] = row[col_indices['description']].strip()
                
                if 'quantity' in col_indices and col_indices['quantity'] < len(row):
                    try:
                        item['quantity'] = float(row[col_indices['quantity']].replace(',', ''))
                    except (ValueError, AttributeError):
                        item['quantity'] = 1.0
                
                if 'rate' in col_indices and col_indices['rate'] < len(row):
                    try:
                        rate_text = row[col_indices['rate']].replace('₹', '').replace(',', '').strip()
                        item['rate'] = float(rate_text)
                    except (ValueError, AttributeError):
                        item['rate'] = 0.0
                
                if 'amount' in col_indices and col_indices['amount'] < len(row):
                    try:
                        amount_text = row[col_indices['amount']].replace('₹', '').replace(',', '').strip()
                        item['amount'] = float(amount_text)
                    except (ValueError, AttributeError):
                        item['amount'] = 0.0
                
                if 'unit' in col_indices and col_indices['unit'] < len(row):
                    item['unit'] = row[col_indices['unit']].strip()
                else:
                    item['unit'] = 'nos'
                
                # Only add if has meaningful content
                if item.get('description') and len(item.get('description', '')) > 3:
                    line_items.append(item)
        
        return line_items
    
    def _extract_line_items_from_dataframes(self, dataframes) -> List[Dict[str, Any]]:
        """Extract line items from pandas dataframes (tabula output)"""
        line_items = []
        
        for df in dataframes:
            if df.empty:
                continue
            
            # Convert dataframe to list of lists
            table_data = [df.columns.tolist()] + df.values.tolist()
            line_items.extend(self._extract_line_items_from_tables([table_data]))
        
        return line_items
    
    def _select_best_result(self, results: List[POExtractedData]) -> POExtractedData:
        """Select the best extraction result based on confidence score and completeness"""
        if not results:
            return POExtractedData(
                raw_text="No extraction results available",
                extraction_method="none",
                confidence_score=0.0
            )
        
        # Sort by confidence score
        results.sort(key=lambda x: x.confidence_score or 0, reverse=True)
        best_result = results[0]
        
        # Try to merge information from other results
        for other_result in results[1:]:
            if not best_result.po_number and other_result.po_number:
                best_result.po_number = other_result.po_number
            if not best_result.vendor_name and other_result.vendor_name:
                best_result.vendor_name = other_result.vendor_name
            if not best_result.total_amount and other_result.total_amount:
                best_result.total_amount = other_result.total_amount
            if not best_result.line_items and other_result.line_items:
                best_result.line_items = other_result.line_items
        
        return best_result


class UserRole(str, Enum):
    SUPER_ADMIN = "super_admin"
    INVOICE_CREATOR = "invoice_creator"
    REVIEWER = "reviewer"
    APPROVER = "approver"
    CLIENT = "client"
    VENDOR = "vendor"

class MasterItem(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    description: str
    unit: str
    standard_rate: float
    category: Optional[str] = None
    last_used_date: Optional[datetime] = None
    usage_count: int = 0
    created_by: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class InvoiceType(str, Enum):
    PROFORMA = "proforma"
    TAX_INVOICE = "tax_invoice"

class InvoiceStatus(str, Enum):
    DRAFT = "draft"
    PENDING_REVIEW = "pending_review"
    REVIEWED = "reviewed"
    APPROVED = "approved"
    PAID = "paid"

# Enhanced Models for Company Profile and Metadata Management

class CompanyLocation(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    location_name: str
    address_line_1: str
    address_line_2: Optional[str] = None
    city: str
    state: str
    pincode: str
    country: str = "India"
    phone: Optional[str] = None
    email: Optional[str] = None
    gst_number: Optional[str] = None
    is_default: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class CompanyBankDetails(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    bank_name: str
    account_number: str
    account_holder_name: str
    ifsc_code: str
    branch_name: str
    account_type: str = "Current"  # Current, Savings
    is_default: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class CompanyProfile(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    company_logo: Optional[str] = None
    locations: List[CompanyLocation] = []
    bank_details: List[CompanyBankDetails] = []
    default_location_id: Optional[str] = None
    default_bank_id: Optional[str] = None
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

# Enhanced Project Metadata Models
class EnhancedProjectMetadata(BaseModel):
    purchase_order_number: str  # Mandatory field
    type: str
    reference_no: Optional[str] = None
    dated: Optional[str] = None
    basic: Optional[float] = 0.0
    overall_multiplier: Optional[float] = 1.0
    po_inv_value: Optional[float] = 0.0
    abg_percentage: Optional[float] = 0.0  # Advance Bank Guarantee %
    ra_bill_with_taxes_percentage: Optional[float] = 0.0  # RA Bill with Taxes %
    erection_percentage: Optional[float] = 0.0  # Erection %
    pbg_percentage: Optional[float] = 0.0  # Performance Bank Guarantee %
    
    # Calculated fields
    abg_amount: Optional[float] = 0.0
    ra_bill_amount: Optional[float] = 0.0
    erection_amount: Optional[float] = 0.0
    pbg_amount: Optional[float] = 0.0
    
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

# Enhanced GST Models
class GSTType(str, Enum):
    CGST_SGST = "cgst_sgst"  # Central + State GST
    IGST = "igst"  # Integrated GST

class ItemGSTMapping(BaseModel):
    item_id: str
    gst_type: GSTType
    cgst_rate: Optional[float] = 0.0
    sgst_rate: Optional[float] = 0.0
    igst_rate: Optional[float] = 0.0
    total_gst_rate: float

# Enhanced RA Bill Tracking Models
class RAQuantityTracking(BaseModel):
    item_id: str
    description: str
    unit: str
    overall_qty: float
    ra_usage: Dict[str, float] = {}  # {"RA1": 10.0, "RA2": 5.0, ...}
    balance_qty: float
    rate: float
    gst_mapping: ItemGSTMapping

class RABillTracking(BaseModel):
    project_id: str
    ra_number: str  # RA1, RA2, RA3, etc.
    items: List[RAQuantityTracking]
    total_amount: float
    gst_amount: float
    total_with_gst: float
    created_at: datetime = Field(default_factory=datetime.utcnow)

# Enhanced Project Model Updates
class EnhancedProject(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_name: str
    architect: str
    client_id: str
    client_name: str
    
    # Enhanced company and location details
    company_profile_id: Optional[str] = None
    selected_location_id: Optional[str] = None
    selected_bank_id: Optional[str] = None
    
    # Project metadata with validation
    metadata: List[EnhancedProjectMetadata] = []
    metadata_validated: bool = False
    validation_errors: List[str] = []
    
    # Enhanced BOQ with RA tracking
    boq_items: List[Dict[str, Any]] = []
    ra_tracking: List[RABillTracking] = []
    
    # Financial calculations
    total_project_value: float = 0.0
    advance_received: float = 0.0
    pending_payment: float = 0.0
    
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

# Enhanced Invoice Model Updates  
class EnhancedInvoice(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    invoice_number: str
    project_id: str
    project_name: str
    client_name: str
    
    # Company and GST details
    company_location_id: Optional[str] = None
    company_bank_id: Optional[str] = None
    invoice_gst_type: GSTType = GSTType.CGST_SGST
    
    # Enhanced invoice items with GST mapping
    invoice_items: List[Dict[str, Any]] = []
    item_gst_mappings: List[ItemGSTMapping] = []
    
    # RA Bill specific fields
    ra_number: Optional[str] = None
    ra_quantities: List[RAQuantityTracking] = []
    quantity_validation_passed: bool = True
    validation_errors: List[str] = []
    
    # Invoice type and amounts
    invoice_type: InvoiceType
    invoice_date: datetime = Field(default_factory=datetime.utcnow)
    
    # Financial details with enhanced GST
    subtotal: float = 0.0
    cgst_amount: float = 0.0
    sgst_amount: float = 0.0
    igst_amount: float = 0.0
    total_gst_amount: float = 0.0
    total_amount: float = 0.0
    
    status: InvoiceStatus = InvoiceStatus.DRAFT
    payment_terms: Optional[str] = None
    advance_received: float = 0.0
    
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    password_hash: str
    role: UserRole
    company_name: Optional[str] = None
    is_active: bool = True
    created_at: datetime = Field(default_factory=datetime.utcnow)

class UserCreate(BaseModel):
    email: str
    password: str
    role: UserRole
    company_name: Optional[str] = None

class UserLogin(BaseModel):
    email: str
    password: str

class ProjectMetadata(BaseModel):
    project_name: Optional[str] = None
    architect: Optional[str] = None
    client: Optional[str] = None
    location: Optional[str] = None
    date: Optional[str] = None

class BOQItem(BaseModel):
    serial_number: str
    description: str
    unit: str
    quantity: float
    rate: float
    amount: float
    category: Optional[str] = None
    billed_quantity: float = 0.0  # Track how much has been billed
    remaining_quantity: Optional[float] = None  # Calculate remaining
    gst_rate: float = 18.0  # GST rate for this item
    is_gst_locked: bool = False  # Lock GST for RA2+ invoices

class InvoiceItem(BaseModel):
    boq_item_id: str  # Reference to original BOQ item
    serial_number: str
    description: str
    unit: str
    quantity: float  # Partial quantity being billed
    rate: float
    amount: float
    gst_rate: float = 18.0
    gst_amount: float = 0.0
    total_with_gst: float = 0.0

class ClientInfo(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    gst_no: Optional[str] = None
    bill_to_address: str
    ship_to_address: Optional[str] = None
    contact_person: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)

class Project(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_name: str
    architect: str
    client_id: Optional[str] = None
    client_name: str
    location: Optional[str] = None
    metadata: Optional[dict] = Field(default_factory=dict)
    boq_items: List[dict] = Field(default_factory=list)
    total_project_value: float = 0.0
    advance_received: float = 0.0
    pending_payment: float = 0.0
    created_by: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class BankGuarantee(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    project_name: str
    guarantee_type: str  # Performance, Advance Payment, Retention, etc.
    guarantee_amount: float
    guarantee_percentage: float
    issuing_bank: str
    guarantee_number: str
    issue_date: datetime
    validity_date: datetime
    beneficiary: str
    applicant: str
    guarantee_details: str
    document_path: Optional[str] = None  # Path to uploaded document
    status: str = "active"  # active, expired, claimed, released
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class Invoice(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    invoice_number: str = ""
    ra_number: str = ""  # RA1, RA2, RA3, etc.
    project_id: str
    project_name: str
    client_id: str
    client_name: str
    invoice_type: InvoiceType
    items: List[InvoiceItem]
    subtotal: float = 0.0
    total_gst_amount: float = 0.0
    total_amount: float = 0.0
    is_partial: bool = True  # Most invoices are partial
    billing_percentage: Optional[float] = None  # What % of project is being billed
    cumulative_billed: Optional[float] = None  # Total billed so far including this invoice
    status: InvoiceStatus = InvoiceStatus.DRAFT
    created_by: Optional[str] = None
    reviewed_by: Optional[str] = None
    approved_by: Optional[str] = None
    invoice_date: datetime = Field(default_factory=datetime.utcnow)
    due_date: Optional[datetime] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class ActivityLog(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    user_email: str
    user_role: str
    action: str
    description: str
    project_id: Optional[str] = None
    invoice_id: Optional[str] = None
    timestamp: datetime = Field(default_factory=datetime.utcnow)

# Excel Parser Class
# GST State mapping for India
GST_STATE_CODES = {
    "andhra pradesh": "AP",
    "arunachal pradesh": "AR", 
    "assam": "AS",
    "bihar": "BR",
    "chhattisgarh": "CG",
    "goa": "GA",
    "gujarat": "GJ",
    "haryana": "HR",
    "himachal pradesh": "HP",
    "jharkhand": "JH",
    "karnataka": "KA",  # Activus base state
    "kerala": "KL",
    "madhya pradesh": "MP",
    "maharashtra": "MH",
    "manipur": "MN",
    "meghalaya": "ML",
    "mizoram": "MZ",
    "nagaland": "NL",
    "odisha": "OR",
    "punjab": "PB",
    "rajasthan": "RJ",
    "sikkim": "SK",
    "tamil nadu": "TN",
    "telangana": "TS",
    "tripura": "TR",
    "uttar pradesh": "UP",
    "uttarakhand": "UK",
    "west bengal": "WB",
    "delhi": "DL",
    "jammu and kashmir": "JK",
    "ladakh": "LA",
    "puducherry": "PY",
    "chandigarh": "CH",
    "andaman and nicobar islands": "AN",
    "dadra and nagar haveli and daman and diu": "DN",
    "lakshadweep": "LD"
}

# Standard GST rates in India
STANDARD_GST_RATES = [5.0, 12.0, 18.0, 28.0]

def determine_gst_type(client_address: str, company_state: str = "karnataka") -> dict:
    """
    Determine GST type (IGST vs CGST+SGST) based on client and company location
    Activus is based in Bangalore, Karnataka
    """
    try:
        # Clean and normalize addresses
        client_address_lower = client_address.lower() if client_address else ""
        company_state_lower = company_state.lower()
        
        # Extract state from client address
        client_state = None
        for state_name, state_code in GST_STATE_CODES.items():
            if state_name in client_address_lower or state_code.lower() in client_address_lower:
                client_state = state_name
                break
        
        # If client state not identified from address, assume different state (IGST)
        if not client_state:
            return {
                "gst_type": "IGST",
                "is_interstate": True,
                "client_state": "Unknown",
                "company_state": "Karnataka"
            }
        
        # Check if same state or different state
        is_same_state = client_state == company_state_lower
        
        return {
            "gst_type": "CGST+SGST" if is_same_state else "IGST",
            "is_interstate": not is_same_state,
            "client_state": client_state.title(),
            "company_state": company_state_lower.title()
        }
        
    except Exception as e:
        logger.error(f"Error determining GST type: {str(e)}")
        # Default to IGST for safety
        return {
            "gst_type": "IGST", 
            "is_interstate": True,
            "client_state": "Unknown",
            "company_state": "Karnataka"
        }

class ExcelParser:
    def __init__(self):
        self.metadata_patterns = {
            'project_name': [r'project\s*name', r'project\s*:', r'job\s*name'],
            'architect': [r'architect', r'architect\s*name', r'architect\s*:'],
            'client': [r'client', r'client\s*name', r'client\s*:'],
            'location': [r'location', r'site', r'address'],
            'date': [r'date', r'project\s*date']
        }
    
    async def parse_excel_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            worksheet = self._select_worksheet(workbook)
            
            if not worksheet:
                raise ValueError("No valid worksheet found in the Excel file")
            
            # Extract metadata
            metadata = self._extract_metadata(worksheet)
            
            # Extract BOQ items
            items = self._extract_boq_items(worksheet)
            
            if not items:
                logger.warning(f"No BOQ items found in file {filename}")
                # Still allow processing with empty items
            
            # Calculate totals with validation
            total_value = 0.0
            try:
                total_value = sum(item['amount'] for item in items if item.get('amount'))
            except (TypeError, ValueError) as e:
                logger.warning(f"Error calculating total value: {e}")
                total_value = 0.0
            
            return {
                'metadata': metadata,
                'items': items,
                'total_value': total_value,
                'filename': filename,
                'items_count': len(items)
            }
        except Exception as e:
            logger.error(f"Error parsing Excel file {filename}: {e}")
            raise HTTPException(
                status_code=400, 
                detail=f"Error parsing Excel file: {str(e)}. Please ensure the file contains BOQ data in the expected format."
            )
    
    def _select_worksheet(self, workbook: openpyxl.Workbook):
        """Select the appropriate worksheet containing BOQ data"""
        # Priority order for worksheet selection
        preferred_names = ['boq', 'bill of quantities', 'quantities', 'estimate', 'summary', 'sheet1', 'main']
        
        # First try to find by preferred names
        for name in preferred_names:
            for sheet_name in workbook.sheetnames:
                if name.lower() in sheet_name.lower():
                    return workbook[sheet_name]
        
        # If no preferred name found, return the first non-empty sheet
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            if worksheet.max_row > 1 and worksheet.max_column > 1:
                return worksheet
        
        # Return active worksheet as last resort
        return workbook.active
    
    def _extract_metadata(self, worksheet) -> Dict[str, Any]:
        metadata = {}
        
        # Search first 20 rows for metadata
        for row_idx in range(1, min(21, worksheet.max_row + 1)):
            for col_idx in range(1, min(10, worksheet.max_column + 1)):
                cell = worksheet.cell(row=row_idx, column=col_idx)
                if cell.value and isinstance(cell.value, str):
                    self._extract_metadata_field(cell.value, worksheet, row_idx, col_idx, metadata)
        
        return metadata
    
    def _extract_metadata_field(self, cell_value: str, worksheet, row_idx: int, col_idx: int, metadata: Dict):
        cell_lower = cell_value.lower().strip()
        
        for field, patterns in self.metadata_patterns.items():
            for pattern in patterns:
                if re.search(pattern, cell_lower):
                    value = self._find_adjacent_value(worksheet, row_idx, col_idx)
                    if value:
                        metadata[field] = value
                        break
    
    def _find_adjacent_value(self, worksheet, row_idx: int, col_idx: int) -> Optional[str]:
        # Check same cell after colon/dash
        current_cell = worksheet.cell(row=row_idx, column=col_idx)
        if current_cell.value and (':' in str(current_cell.value) or '-' in str(current_cell.value)):
            parts = re.split(r'[:\-]', str(current_cell.value), 1)
            if len(parts) > 1 and parts[1].strip():
                return parts[1].strip()
        
        # Check right cell
        right_cell = worksheet.cell(row=row_idx, column=col_idx + 1)
        if right_cell.value and str(right_cell.value).strip():
            return str(right_cell.value).strip()
        
        return None
    
    def _extract_boq_items(self, worksheet) -> List[Dict]:
        items = []
        header_row = self._find_header_row(worksheet)
        
        if not header_row:
            return items
        
        column_mapping = self._map_columns(worksheet, header_row)
        serial_number = 1
        
        for row_idx in range(header_row + 1, worksheet.max_row + 1):
            row_data = self._extract_row_data(worksheet, row_idx, column_mapping)
            
            if self._is_valid_item_row(row_data):
                # Extract description using safe string conversion
                description = self._safe_string_conversion(row_data.get('description'))
                if len(description) < 3:  # Skip very short descriptions
                    continue
                
                # Extract other fields with better handling
                unit = "nos"  # default unit
                unit_value = row_data.get('unit')
                if unit_value and str(unit_value).strip():
                    unit_text = str(unit_value).strip()
                    # Handle common unit patterns
                    if unit_text.lower() in ['cum', 'cu.m', 'cubic meter', 'cubicmeter']:
                        unit = 'Cum'
                    elif unit_text.lower() in ['sqm', 'sq.m', 'square meter', 'squaremeter']:
                        unit = 'Sqm'
                    elif unit_text.lower() in ['rmt', 'rm', 'running meter']:
                        unit = 'Rmt'
                    elif unit_text.lower() in ['nos', 'no', 'number', 'each']:
                        unit = 'Nos'
                    elif unit_text.lower() in ['kg', 'kilogram']:
                        unit = 'Kg'
                    elif unit_text.lower() in ['ton', 'tonne', 'mt']:
                        unit = 'Ton'
                    elif unit_text.lower() in ['ltr', 'litre', 'liter']:
                        unit = 'Ltr'
                    else:
                        # If it's not a common numeric value, keep as is
                        try:
                            float_val = float(unit_text)
                            # If it converts to float successfully and is a large number, it's probably wrong data
                            if float_val > 100:
                                unit = 'Nos'  # fallback
                            else:
                                unit = unit_text
                        except ValueError:
                            # If it can't be converted to float, it's likely the correct unit text
                            unit = unit_text
                
                quantity = self._safe_float_conversion(row_data.get('quantity'))
                if quantity == 0.0:
                    quantity = 1.0  # default quantity
                
                rate = self._safe_float_conversion(row_data.get('rate'))
                
                amount = self._safe_float_conversion(row_data.get('amount'))
                if amount == 0.0:
                    amount = quantity * rate if rate > 0 else 0.0
                
                items.append({
                    'serial_number': serial_number,
                    'description': description,
                    'unit': unit,
                    'quantity': quantity,
                    'rate': rate,
                    'amount': amount,
                    'gst_rate': 18.0  # Default GST rate
                })
                
                serial_number += 1
        
        return items
    
    def _find_header_row(self, worksheet) -> Optional[int]:
        # Look for rows that contain typical BOQ headers
        header_keywords = ['description', 'quantity', 'rate', 'amount', 'item', 'particular', 'unit', 'uom']
        
        for row_idx in range(1, min(30, worksheet.max_row + 1)):
            row_cells = []
            for col in range(1, min(20, worksheet.max_column + 1)):
                cell_value = worksheet.cell(row=row_idx, column=col).value
                if cell_value:
                    row_cells.append(str(cell_value).lower().strip())
            
            row_text = ' '.join(row_cells)
            print(f"Checking row {row_idx}: {row_text[:100]}...")
            
            # Count matches and ensure we have essential columns
            matches = sum(1 for keyword in header_keywords if keyword in row_text)
            has_description = any(desc in row_text for desc in ['description', 'item', 'particular', 'work'])
            has_quantity = any(qty in row_text for qty in ['quantity', 'qty'])
            has_rate = any(rate in row_text for rate in ['rate', 'price'])
            
            print(f"Row {row_idx} matches: {matches}, has_desc: {has_description}, has_qty: {has_quantity}, has_rate: {has_rate}")
            
            if matches >= 3 and has_description and (has_quantity or has_rate):
                print(f"Selected header row: {row_idx}")
                return row_idx
        
        print("No suitable header row found, using fallback")
        return None
    
    def _map_columns(self, worksheet, header_row: int) -> Dict[str, int]:
        column_mapping = {}
        
        print(f"Mapping columns from header row {header_row}")
        
        # First pass: collect all headers and their positions
        headers = []
        for col_idx in range(1, worksheet.max_column + 1):
            cell_value = worksheet.cell(row=header_row, column=col_idx).value
            if cell_value:
                cell_lower = str(cell_value).lower().strip()
                headers.append((col_idx, cell_lower, str(cell_value)))
                print(f"Column {col_idx}: '{cell_value}' -> '{cell_lower}'")
        
        # Second pass: map columns with priority and exclusions
        for col_idx, cell_lower, original_value in headers:
            
            # Serial Number - must be first
            if col_idx <= 2 and any(h in cell_lower for h in ['s.no', 'sr.no', 'serial', 'sl', 'sno']):
                column_mapping['serial'] = col_idx
                print(f"Serial column mapped to {col_idx}")
                
            # Description - usually first text column after serial
            elif any(h in cell_lower for h in ['description', 'item', 'particular', 'work', 'scope']) and 'unit' not in cell_lower and 'rate' not in cell_lower:
                column_mapping['description'] = col_idx
                print(f"Description column mapped to {col_idx}")
                
            # Unit - specific keywords, exclude rate-related terms
            elif any(h in cell_lower for h in ['unit', 'uom', 'u.o.m']) and 'rate' not in cell_lower and 'price' not in cell_lower and 'amount' not in cell_lower:
                column_mapping['unit'] = col_idx
                print(f"Unit column mapped to {col_idx}")
                
            # Quantity - before rate column
            elif any(h in cell_lower for h in ['qty', 'quantity']) and 'rate' not in cell_lower and 'unit' not in cell_lower:
                column_mapping['quantity'] = col_idx
                print(f"Quantity column mapped to {col_idx}")
                
            # Rate - must have rate/price but not be amount/total
            elif any(h in cell_lower for h in ['rate', 'price']) and 'amount' not in cell_lower and 'total' not in cell_lower:
                if 'unit' in cell_lower:
                    column_mapping['rate'] = col_idx
                    print(f"Rate column mapped to {col_idx} (unit rate)")
                elif not column_mapping.get('rate'):  # Only if rate not already mapped
                    column_mapping['rate'] = col_idx
                    print(f"Rate column mapped to {col_idx}")
                    
            # Amount - must have amount/total but not be rate
            elif any(h in cell_lower for h in ['amount', 'total']) and 'rate' not in cell_lower and 'unit' not in cell_lower:
                column_mapping['amount'] = col_idx
                print(f"Amount column mapped to {col_idx}")
        
        # Validation: ensure we have essential columns
        if not column_mapping.get('description'):
            # Fallback: find first text-heavy column
            for col_idx in range(1, min(6, worksheet.max_column + 1)):
                if col_idx not in column_mapping.values():
                    column_mapping['description'] = col_idx
                    print(f"Fallback: Description mapped to column {col_idx}")
                    break
                    
        if not column_mapping.get('unit'):
            # Look for column that typically has text values, not numbers
            for col_idx in range(2, min(7, worksheet.max_column + 1)):
                if col_idx not in column_mapping.values():
                    # Check if this column has text values in the next few rows
                    text_count = 0
                    for row in range(header_row + 1, min(header_row + 6, worksheet.max_row + 1)):
                        cell_val = worksheet.cell(row=row, column=col_idx).value
                        if cell_val and isinstance(cell_val, str) and not cell_val.replace('.', '').isdigit():
                            text_count += 1
                    if text_count > 0:
                        column_mapping['unit'] = col_idx
                        print(f"Fallback: Unit mapped to column {col_idx} based on text content")
                        break
        
        print(f"Final column mapping: {column_mapping}")
        return column_mapping
    
    def _extract_row_data(self, worksheet, row_idx: int, column_mapping: Dict[str, int]) -> Dict:
        row_data = {}
        
        for field, col_idx in column_mapping.items():
            cell = worksheet.cell(row=row_idx, column=col_idx)
            cell_value = cell.value
            
            print(f"Row {row_idx}, Column {col_idx} ({field}): '{cell_value}' (type: {type(cell_value)})")
            
            # Special handling for different field types
            if field == 'unit':
                if cell_value is not None:
                    # Convert to string and clean up
                    unit_str = str(cell_value).strip()
                    # Remove any trailing decimals from numeric strings that should be text
                    if unit_str.endswith('.0'):
                        unit_str = unit_str[:-2]
                    # Check if it looks like a unit (contains letters or common unit abbreviations)
                    if any(c.isalpha() for c in unit_str) or unit_str.lower() in ['cum', 'sqm', 'nos', 'rmt', 'kg', 'ltr']:
                        row_data[field] = unit_str
                    else:
                        # If it's purely numeric and doesn't look like a unit, set default
                        row_data[field] = 'Nos'
                        print(f"Warning: Column {col_idx} has numeric value '{unit_str}' for unit, using default 'Nos'")
                else:
                    row_data[field] = 'Nos'  # default unit
                    
            elif field in ['quantity', 'rate', 'amount']:
                # Ensure numeric fields are properly converted
                row_data[field] = self._safe_float_conversion(cell_value)
                
            else:
                row_data[field] = cell_value
        
        return row_data
    
    def _is_valid_item_row(self, row_data: Dict) -> bool:
        description = row_data.get('description', '')
        if not description or not isinstance(description, str) or len(description.strip()) < 3:
            return False
        
        # Must have at least description and one numeric field with a value > 0
        numeric_fields = ['quantity', 'rate', 'amount']
        has_numeric = any(
            self._safe_float_conversion(row_data.get(field)) > 0
            for field in numeric_fields
        )
        
        return has_numeric
    
    def _safe_float_conversion(self, value):
        """Safely convert various data types to float, handling currency symbols and empty values"""
        if value is None or value == "":
            return 0.0
        
        if isinstance(value, (int, float)):
            return float(value)
        
        if isinstance(value, str):
            # Remove common currency symbols and whitespace
            cleaned_value = str(value).replace('₹', '').replace('Rs', '').replace(',', '').strip()
            if cleaned_value == "" or cleaned_value.lower() == "none":
                return 0.0
            try:
                return float(cleaned_value)
            except (ValueError, TypeError):
                return 0.0
        
        try:
            return float(value)
        except (ValueError, TypeError):
            return 0.0

    def _safe_string_conversion(self, value):
        """Safely convert value to string, handling None and empty values"""
        if value is None or value == "":
            return ""
        return str(value).strip()

# PDF Generator Class
class PDFGenerator:
    def __init__(self):
        self.page_size = A4
        self.margin = 20 * mm
        
    async def generate_invoice_pdf(self, invoice: Invoice, project: Project, client: ClientInfo) -> BytesIO:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=self.page_size,
            rightMargin=self.margin,
            leftMargin=self.margin,
            topMargin=self.margin,
            bottomMargin=self.margin
        )
        
        # Build PDF content
        elements = []
        styles = getSampleStyleSheet()
        
        # Company brand color
        company_color = colors.HexColor('#127285')
        light_bg_color = colors.HexColor('#f8f9fa')
        
        # Add custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=company_color,
            alignment=TA_CENTER,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#666666'),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=company_color,
            spaceAfter=10,
            fontName='Helvetica-Bold'
        )
        
        # Company Header with Logo Space
        elements.append(Paragraph("ACTIVUS INDUSTRIAL DESIGN & BUILD LLP", title_style))
        elements.append(Paragraph("One Stop Solution for Industrial Projects", subtitle_style))
        
        # Invoice Title
        invoice_type_display = "TAX INVOICE" if invoice.invoice_type.value == "tax_invoice" else "PROFORMA INVOICE"
        elements.append(Paragraph(invoice_type_display, header_style))
        elements.append(Spacer(1, 20))
        
        # Invoice Details Table
        details_data = [
            ['Invoice Number:', invoice.invoice_number, 'Date:', invoice.invoice_date.strftime('%d/%m/%Y')],
            ['Project:', project.project_name, 'Client:', client.name],
            ['Architect:', project.architect, '', '']
        ]
        
        details_table = Table(details_data, colWidths=[80, 200, 60, 120])
        details_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), light_bg_color),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dddddd')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(details_table)
        elements.append(Spacer(1, 20))
        
        # Bill To Section
        elements.append(Paragraph("Bill To:", header_style))
        bill_to_text = f"""
        <b>{client.name}</b><br/>
        {client.bill_to_address}<br/>
        <b>GST No:</b> {client.gst_no if client.gst_no else 'Not Available'}
        """
        elements.append(Paragraph(bill_to_text, styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Items Table with proper text wrapping
        table_data = [
            ['S.No', 'Description', 'Unit', 'Qty', 'Rate (Rs)', 'Amount (Rs)']
        ]
        
        for idx, item in enumerate(invoice.items, 1):
            # Format currency properly - use Rs instead of ₹ symbol
            rate_formatted = f"Rs {item.rate:,.2f}"
            amount_formatted = f"Rs {item.amount:,.2f}"
            
            # Wrap long descriptions
            description_wrapped = item.description
            if len(description_wrapped) > 60:
                # Break into multiple lines for long descriptions
                words = description_wrapped.split(' ')
                lines = []
                current_line = []
                current_length = 0
                
                for word in words:
                    if current_length + len(word) > 60:
                        if current_line:
                            lines.append(' '.join(current_line))
                            current_line = [word]
                            current_length = len(word)
                        else:
                            lines.append(word)
                            current_length = 0
                    else:
                        current_line.append(word)
                        current_length += len(word) + 1
                
                if current_line:
                    lines.append(' '.join(current_line))
                
                description_wrapped = '<br/>'.join(lines)
            
            table_data.append([
                str(idx),
                Paragraph(description_wrapped, styles['Normal']),
                item.unit,
                f"{item.quantity:,.1f}",
                rate_formatted,
                amount_formatted
            ])
        
        # Create table with fixed column widths
        col_widths = [30, 240, 50, 50, 80, 90]  # Total: 540
        items_table = Table(table_data, colWidths=col_widths)
        
        # Enhanced table styling
        items_table.setStyle(TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), company_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            
            # Data rows styling
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # S.No center
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),    # Description left
            ('ALIGN', (2, 1), (2, -1), 'CENTER'),  # Unit center
            ('ALIGN', (3, 1), (3, -1), 'CENTER'),  # Qty center
            ('ALIGN', (4, 1), (-1, -1), 'RIGHT'),  # Rate and Amount right
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            
            # Grid and padding
            ('GRID', (0, 0), (-1, -1), 1.5, colors.HexColor('#cccccc')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            
            # Special styling for description column
            ('LEFTPADDING', (1, 1), (1, -1), 6),
            ('RIGHTPADDING', (1, 1), (1, -1), 6),
        ]))
        
        elements.append(items_table)
        elements.append(Spacer(1, 20))
        
        # Determine GST breakdown
        gst_info = invoice.gst_info if hasattr(invoice, 'gst_info') else {"gst_type": "IGST"}
        
        # Totals Table with proper GST display
        if gst_info.get("gst_type") == "CGST+SGST":
            totals_data = [
                ['', '', '', '', 'Subtotal:', f"Rs {invoice.subtotal:,.2f}"],
                ['', '', '', '', 'CGST (9%):', f"Rs {invoice.total_gst_amount/2:,.2f}"],
                ['', '', '', '', 'SGST (9%):', f"Rs {invoice.total_gst_amount/2:,.2f}"],
                ['', '', '', '', 'Total Amount:', f"Rs {invoice.total_amount:,.2f}"]
            ]
        else:
            # IGST
            gst_rate = 18  # Default, could be calculated from items
            totals_data = [
                ['', '', '', '', 'Subtotal:', f"Rs {invoice.subtotal:,.2f}"],
                ['', '', '', '', f'IGST ({gst_rate}%):', f"Rs {invoice.total_gst_amount:,.2f}"],
                ['', '', '', '', 'Total Amount:', f"Rs {invoice.total_amount:,.2f}"]
            ]
        
        totals_table = Table(totals_data, colWidths=col_widths)
        totals_table.setStyle(TableStyle([
            ('ALIGN', (4, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (4, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (4, 0), (-1, -1), 11),
            ('TEXTCOLOR', (4, 0), (-1, -1), colors.black),
            ('BOX', (4, 0), (-1, -1), 1.5, colors.HexColor('#cccccc')),
            ('INNERGRID', (4, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
            ('BACKGROUND', (4, -1), (-1, -1), colors.HexColor('#e8f4f8')),  # Highlight total row
            ('LEFTPADDING', (4, 0), (-1, -1), 8),
            ('RIGHTPADDING', (4, 0), (-1, -1), 8),
            ('TOPPADDING', (4, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (4, 0), (-1, -1), 8),
        ]))
        
        elements.append(totals_table)
        elements.append(Spacer(1, 30))
        
        # Terms and Conditions
        terms_text = """
        <b>Terms & Conditions:</b><br/>
        1. Payment to be made within 30 days from the date of invoice.<br/>
        2. All disputes subject to Bangalore jurisdiction.<br/>
        3. Goods once sold will not be taken back.<br/>
        """
        elements.append(Paragraph(terms_text, styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Footer
        footer_text = """
        <para alignment="center">
        <b>Thank you for your business!</b><br/>
        <b>ACTIVUS INDUSTRIAL DESIGN & BUILD LLP</b><br/>
        For any queries, please contact us at info@activusdesign.com | +91-XXXXXXXXXX
        </para>
        """
        elements.append(Paragraph(footer_text, styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer

# Authentication functions
async def hash_password(password: str) -> str:
    salt = bcrypt.gensalt()
    return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')

async def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

async def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        'user_id': user_id,
        'email': email,
        'role': role,
        'exp': datetime.utcnow() + timedelta(days=7)
    }
    return jwt.encode(payload, SECRET_KEY, algorithm='HS256')

async def verify_token(token: str) -> Dict:
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    payload = await verify_token(credentials.credentials)
    user = await db.users.find_one({"id": payload["user_id"]})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

async def log_activity(user_id: str, user_email: str, user_role: str, action: str, description: str, 
                      project_id: Optional[str] = None, invoice_id: Optional[str] = None):
    log_entry = ActivityLog(
        user_id=user_id,
        user_email=user_email,
        user_role=user_role,
        action=action,
        description=description,
        project_id=project_id,
        invoice_id=invoice_id
    )
    await db.activity_logs.insert_one(log_entry.dict())

# Initialize super admin
async def init_super_admin():
    super_admin_email = "brightboxm@gmail.com"
    existing_user = await db.users.find_one({"email": super_admin_email})
    
    if not existing_user:
        password_hash = await hash_password("admin123")
        super_admin = User(
            email=super_admin_email,
            password_hash=password_hash,
            role=UserRole.SUPER_ADMIN,
            company_name="Activus Industrial Design & Build"
        )
        await db.users.insert_one(super_admin.dict())
        logger.info("Super admin created successfully")

# API Routes
@api_router.post("/auth/login")
async def login(user_data: UserLogin):
    try:
        user = await db.users.find_one({"email": user_data.email, "is_active": True})
        if not user:
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        if not await verify_password(user_data.password, user["password_hash"]):
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        token = await create_token(user["id"], user["email"], user["role"])
        
        await log_activity(
            user["id"], user["email"], user["role"], 
            "login", f"User logged in successfully"
        )
        
        return {
            "access_token": token,
            "token_type": "bearer",
            "user": {
                "id": user["id"],
                "email": user["email"],
                "role": user["role"],
                "company_name": user["company_name"]
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Login error: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@api_router.post("/auth/register")
async def register(user_data: UserCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can create users")
    
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    password_hash = await hash_password(user_data.password)
    new_user = User(
        email=user_data.email,
        password_hash=password_hash,
        role=user_data.role,
        company_name=user_data.company_name
    )
    
    await db.users.insert_one(new_user.dict())
    
    await log_activity(
        current_user["id"], current_user["email"], current_user["role"],
        "user_created", f"Created new user: {user_data.email} with role: {user_data.role}"
    )
    
    return {"message": "User created successfully", "user_id": new_user.id}

# User Management APIs
@api_router.get("/users", response_model=List[dict])
async def get_users(current_user: dict = Depends(get_current_user)):
    """Get all users (super admin only)"""
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can view users")
    
    try:
        users = await db.users.find().to_list(1000)
        
        # Clean user data (remove password hash)
        clean_users = []
        for user in users:
            clean_user = {
                "id": user.get("id"),
                "email": user.get("email"),
                "role": user.get("role"),
                "company_name": user.get("company_name"),
                "is_active": user.get("is_active", True),
                "created_at": user.get("created_at"),
                "last_login": user.get("last_login")
            }
            clean_users.append(clean_user)
        
        return clean_users
        
    except Exception as e:
        logger.error(f"Error fetching users: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch users: {str(e)}")

@api_router.put("/users/{user_id}", response_model=dict)
async def update_user(
    user_id: str,
    update_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update user details (super admin only)"""
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can update users")
    
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")
        
        user = await db.users.find_one({"id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Only allow updating specific fields
        allowed_fields = ["role", "company_name", "is_active"]
        filtered_data = {k: v for k, v in update_data.items() if k in allowed_fields}
        
        if not filtered_data:
            raise HTTPException(status_code=400, detail="No valid fields to update")
        
        filtered_data["updated_at"] = datetime.utcnow()
        
        await db.users.update_one(
            {"id": user_id},
            {"$set": filtered_data}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "user_updated", f"Updated user: {user['email']}"
        )
        
        return {"message": "User updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating user: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update user: {str(e)}")

@api_router.delete("/users/{user_id}")
async def deactivate_user(user_id: str, current_user: dict = Depends(get_current_user)):
    """Deactivate a user (super admin only)"""
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can deactivate users")
    
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")
        
        # Prevent super admin from deactivating themselves
        if user_id == current_user["id"]:
            raise HTTPException(status_code=400, detail="Cannot deactivate your own account")
        
        user = await db.users.find_one({"id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        await db.users.update_one(
            {"id": user_id},
            {"$set": {"is_active": False, "updated_at": datetime.utcnow()}}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "user_deactivated", f"Deactivated user: {user['email']}"
        )
        
        return {"message": "User deactivated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deactivating user: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to deactivate user: {str(e)}")

@api_router.post("/users/{user_id}/reset-password")
async def reset_user_password(
    user_id: str,
    password_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Reset user password (super admin only)"""
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can reset passwords")
    
    try:
        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")
        
        new_password = password_data.get("new_password")
        if not new_password or len(new_password) < 6:
            raise HTTPException(status_code=400, detail="Password must be at least 6 characters long")
        
        user = await db.users.find_one({"id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        password_hash = await hash_password(new_password)
        
        await db.users.update_one(
            {"id": user_id},
            {"$set": {"password_hash": password_hash, "updated_at": datetime.utcnow()}}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "password_reset", f"Reset password for user: {user['email']}"
        )
        
        return {"message": "Password reset successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error resetting password: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to reset password: {str(e)}")

@api_router.post("/upload-boq")
async def upload_boq(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    # Validate file type
    allowed_content_types = [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
        'application/vnd.ms-excel.sheet.macroEnabled.12'
    ]
    
    allowed_extensions = ['.xlsx', '.xlsm', '.xls']
    
    if file.content_type not in allowed_content_types:
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid file type. Please upload an Excel file (.xlsx, .xls, .xlsm). Received: {file.content_type}"
        )
    
    if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid file extension. Please upload an Excel file with .xlsx, .xls, or .xlsm extension"
        )
    
    # Check file size (max 10MB)
    file_content = await file.read()
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB")
    
    if len(file_content) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded")
    
    try:
        parser = ExcelParser()
        parsed_data = await parser.parse_excel_file(file_content, file.filename)
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "boq_uploaded", f"Successfully uploaded and parsed BOQ file: {file.filename}"
        )
        
        return parsed_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"BOQ parsing error for file {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=422, 
            detail=f"Failed to parse Excel file. Please ensure it's a valid BOQ format. Error: {str(e)}"
        )

@api_router.post("/clients", response_model=dict)
async def create_client(client_data: ClientInfo, current_user: dict = Depends(get_current_user)):
    await db.clients.insert_one(client_data.dict())
    
    await log_activity(
        current_user["id"], current_user["email"], current_user["role"],
        "client_created", f"Created client: {client_data.name}"
    )
    
    return {"message": "Client created successfully", "client_id": client_data.id}

@api_router.get("/clients", response_model=List[ClientInfo])
async def get_clients(current_user: dict = Depends(get_current_user)):
    clients = await db.clients.find().to_list(1000)
    return [ClientInfo(**client) for client in clients]

@api_router.post("/projects", response_model=dict)
async def create_project(project_data: Project, current_user: dict = Depends(get_current_user)):
    try:
        # Set additional fields
        project_data.created_by = current_user["id"]
        project_data.pending_payment = project_data.total_project_value - project_data.advance_received
        project_data.updated_at = datetime.utcnow()
        
        # Validate BOQ items are properly formed
        if not project_data.boq_items:
            raise HTTPException(status_code=400, detail="BOQ items cannot be empty")
        
        # Insert into database
        await db.projects.insert_one(project_data.dict())
        
        # Log activity
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "project_created", 
            f"Created project: {project_data.project_name} (Value: ₹{project_data.total_project_value:,.2f}, Advance: ₹{project_data.advance_received:,.2f})",
            project_id=project_data.id
        )
        
        return {"message": "Project created successfully", "project_id": project_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating project: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create project: {str(e)}")

@api_router.get("/projects", response_model=List[Project])
async def get_projects(current_user: dict = Depends(get_current_user)):
    try:
        projects = await db.projects.find().to_list(1000)
        
        # Filter and validate projects to prevent null errors
        valid_projects = []
        for project in projects:
            if not project or not isinstance(project, dict):
                continue
                
            # Ensure required fields exist with defaults
            cleaned_project = {
                "id": project.get("id", str(uuid.uuid4())),
                "project_name": project.get("project_name", "Untitled Project"),
                "architect": project.get("architect", "Unknown Architect"),
                "client_id": project.get("client_id", ""),
                "client_name": project.get("client_name", "Unknown Client"),
                "metadata": project.get("metadata", {}),
                "boq_items": project.get("boq_items", []),
                "total_project_value": float(project.get("total_project_value", 0)),
                "advance_received": float(project.get("advance_received", 0)),
                "pending_payment": float(project.get("pending_payment", 0)),
                "created_by": project.get("created_by"),
                "created_at": project.get("created_at", datetime.utcnow()),
                "updated_at": project.get("updated_at", datetime.utcnow())
            }
            
            try:
                valid_project = Project(**cleaned_project)
                valid_projects.append(valid_project)
            except Exception as e:
                logger.warning(f"Skipping invalid project {project.get('id', 'unknown')}: {e}")
                continue
        
        return valid_projects
        
    except Exception as e:
        logger.error(f"Error fetching projects: {str(e)}")
        return []

@api_router.get("/projects/{project_id}", response_model=Project)
async def get_project(project_id: str, current_user: dict = Depends(get_current_user)):
    try:
        if not project_id or len(project_id.strip()) == 0:
            raise HTTPException(status_code=400, detail="Project ID is required")
        
        project = await db.projects.find_one({"id": project_id})
        if not project:
            raise HTTPException(status_code=404, detail=f"Project with ID {project_id} not found")
        
        return Project(**project)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving project {project_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@api_router.get("/projects/{project_id}/details")
async def get_project_details(project_id: str, current_user: dict = Depends(get_current_user)):
    """Get detailed project information with financial summary and related invoices"""
    try:
        # Get project
        project = await db.projects.find_one({"id": project_id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        # Get related invoices
        invoices = await db.invoices.find({"project_id": project_id}).sort("created_at", -1).to_list(1000)
        
        # Calculate financial summary
        total_project_value = project.get("total_project_value", 0)
        total_invoiced = sum(invoice.get("total_amount", 0) for invoice in invoices)
        balance_value = total_project_value - total_invoiced
        
        # Get client details
        client = await db.clients.find_one({"id": project.get("client_id")}) if project.get("client_id") else None
        
        # Clean up invoice data for frontend
        cleaned_invoices = []
        for invoice in invoices:
            cleaned_invoice = {
                "id": invoice.get("id"),
                "invoice_number": invoice.get("invoice_number"),
                "ra_number": invoice.get("ra_number"), 
                "invoice_type": invoice.get("invoice_type"),
                "total_amount": invoice.get("total_amount", 0),
                "subtotal": invoice.get("subtotal", 0),
                "total_gst_amount": invoice.get("total_gst_amount", 0),
                "status": invoice.get("status", "draft"),
                "invoice_date": invoice.get("invoice_date"),
                "created_at": invoice.get("created_at"),
                "gst_info": invoice.get("gst_info", {})
            }
            cleaned_invoices.append(cleaned_invoice)
        
        # Prepare detailed response
        project_details = {
            "project_info": {
                "id": project.get("id"),
                "project_name": project.get("project_name"),
                "architect": project.get("architect"),
                "location": project.get("location"),
                "created_at": project.get("created_at"),
                "updated_at": project.get("updated_at")
            },
            "client_info": {
                "id": client.get("id") if client else None,
                "name": client.get("name") if client else project.get("client_name", "Unknown Client"),
                "bill_to_address": client.get("bill_to_address") if client else "",
                "gst_no": client.get("gst_no") if client else ""
            },
            "financial_summary": {
                "total_project_value": total_project_value,
                "total_invoiced": total_invoiced,
                "balance_value": balance_value,
                "percentage_billed": (total_invoiced / total_project_value * 100) if total_project_value > 0 else 0,
                "total_invoices": len(invoices)
            },
            "boq_summary": {
                "total_items": len(project.get("boq_items", [])),
                "items": project.get("boq_items", [])
            },
            "invoices": cleaned_invoices
        }
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "project_details_viewed", f"Viewed detailed information for project: {project.get('project_name')}"
        )
        
        return project_details
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting project details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get project details: {str(e)}")

@api_router.get("/projects/{project_id}/boq-status")
async def get_project_boq_status(project_id: str, current_user: dict = Depends(get_current_user)):
    """Get BOQ items with billing status for partial invoicing"""
    try:
        project = await db.projects.find_one({"id": project_id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        # Get all invoices for this project
        invoices = await db.invoices.find({"project_id": project_id}).to_list(1000)
        
        # Calculate billing status for each BOQ item
        boq_items_with_status = []
        for boq_item in project.get("boq_items", []):
            item_id = boq_item.get("id", boq_item.get("serial_number"))
            total_billed = 0.0
            
            # Calculate total billed quantity for this item
            for invoice in invoices:
                for inv_item in invoice.get("items", []):
                    if inv_item.get("boq_item_id") == item_id:
                        total_billed += inv_item.get("quantity", 0)
            
            original_quantity = boq_item.get("quantity", 0)
            remaining_quantity = max(0, original_quantity - total_billed)
            billing_percentage = (total_billed / original_quantity * 100) if original_quantity > 0 else 0
            
            boq_item_status = {
                **boq_item,
                "billed_quantity": total_billed,
                "remaining_quantity": remaining_quantity,
                "billing_percentage": round(billing_percentage, 2),
                "is_fully_billed": remaining_quantity == 0,
                "can_bill": remaining_quantity > 0
            }
            
            boq_items_with_status.append(boq_item_status)
        
        # Calculate project-level billing status
        project_total_value = project.get("total_project_value", 0)
        total_billed_value = sum(invoice.get("subtotal", 0) for invoice in invoices)
        project_billing_percentage = (total_billed_value / project_total_value * 100) if project_total_value > 0 else 0
        
        return {
            "project_id": project_id,
            "project_name": project.get("project_name"),
            "total_project_value": project_total_value,
            "total_billed_value": total_billed_value,
            "remaining_value": project_total_value - total_billed_value,
            "project_billing_percentage": round(project_billing_percentage, 2),
            "total_invoices": len(invoices),
            "next_ra_number": f"RA{len(invoices) + 1}",
            "boq_items": boq_items_with_status
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting BOQ status: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get BOQ status")

@api_router.post("/invoices", response_model=dict)
async def create_invoice(invoice_data: dict, current_user: dict = Depends(get_current_user)):
    try:
        # Get project and client information
        project = await db.projects.find_one({"id": invoice_data["project_id"]})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        client = await db.clients.find_one({"id": invoice_data["client_id"]})
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
        
        # Determine GST type based on client location
        gst_info = determine_gst_type(client.get("bill_to_address", ""))
        
        # Check if proforma without tax option
        include_tax = invoice_data.get("include_tax", True)
        is_proforma = invoice_data["invoice_type"] == "proforma"
        
        # Process invoice items
        processed_items = []
        basic_total = 0.0
        total_gst_amount = 0.0
        
        for item_data in invoice_data["items"]:
            # Validate GST rate
            gst_rate = float(item_data.get("gst_rate", 18.0))
            if gst_rate not in STANDARD_GST_RATES:
                gst_rate = 18.0  # Default to 18% if invalid rate
            
            quantity = float(item_data["quantity"])
            rate = float(item_data["rate"])
            amount = quantity * rate
            
            # Calculate GST (zero for proforma without tax)
            if include_tax:
                gst_amount = (amount * gst_rate) / 100
            else:
                gst_amount = 0.0
                gst_rate = 0.0  # Set to 0 for display purposes
            
            total_with_gst = amount + gst_amount
            
            processed_item = {
                "boq_item_id": item_data.get("boq_item_id", item_data.get("serial_number", str(len(processed_items) + 1))),
                "serial_number": len(processed_items) + 1,
                "description": item_data["description"],
                "unit": item_data["unit"],
                "quantity": quantity,
                "rate": rate,
                "amount": amount,
                "gst_rate": gst_rate,
                "gst_amount": gst_amount,
                "total_with_gst": total_with_gst,
                "gst_type": gst_info["gst_type"] if include_tax else "No Tax"
            }
            
            processed_items.append(processed_item)
            basic_total += amount
            total_gst_amount += gst_amount
        
        # Calculate totals
        grand_total = basic_total + total_gst_amount
        
        # Generate invoice and RA numbers
        invoice_count = await db.invoices.count_documents({}) + 1
        invoice_number = f"INV-{datetime.now().year}-{invoice_count:04d}"
        
        # RA numbers only for tax invoices
        ra_number = ""
        if invoice_data["invoice_type"] == "tax_invoice":
            # Count existing tax invoices to get next RA number
            tax_invoice_count = await db.invoices.count_documents({"invoice_type": "tax_invoice"}) + 1
            ra_number = f"RA{tax_invoice_count}"
        
        # Payment terms
        payment_terms = invoice_data.get("payment_terms", "Payment due within 30 days from invoice date")
        
        # Advance against invoice
        advance_received_invoice = float(invoice_data.get("advance_received", 0))
        net_amount = grand_total - advance_received_invoice
        
        # Create invoice
        new_invoice = {
            "id": str(uuid.uuid4()),
            "invoice_number": invoice_number,
            "ra_number": ra_number,
            "project_id": invoice_data["project_id"],
            "project_name": project["project_name"],
            "client_id": invoice_data["client_id"],
            "client_name": client["name"],
            "invoice_type": invoice_data["invoice_type"],
            "include_tax": include_tax,
            "items": processed_items,
            "subtotal": basic_total,  # Basic total before GST
            "total_gst_amount": total_gst_amount,  # Total GST
            "total_amount": grand_total,  # Grand total
            "advance_received": advance_received_invoice,  # Advance against this invoice
            "net_amount": net_amount,  # Amount due after advance
            "payment_terms": payment_terms,
            "gst_info": gst_info,
            "is_partial": invoice_data.get("is_partial", True),
            "billing_percentage": invoice_data.get("billing_percentage"),
            "cumulative_billed": invoice_data.get("cumulative_billed"),
            "status": "draft",
            "created_by": current_user["id"],
            "invoice_date": datetime.utcnow(),
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        # Save to database
        await db.invoices.insert_one(new_invoice)
        
        # Update project advance if advance received against invoice
        if advance_received_invoice > 0:
            await db.projects.update_one(
                {"id": invoice_data["project_id"]},
                {"$inc": {"advance_received": advance_received_invoice}}
            )
        
        # Log activity
        tax_status = "with Tax" if include_tax else "without Tax"
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "invoice_created", 
            f"Created {invoice_data['invoice_type']} invoice {invoice_number} ({tax_status}) for project {project['project_name']} - ₹{grand_total:,.2f}"
        )
        
        return {
            "message": "Invoice created successfully", 
            "invoice_id": new_invoice["id"], 
            "invoice_number": invoice_number,
            "total_amount": grand_total,
            "net_amount": net_amount
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating invoice: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create invoice: {str(e)}")

@api_router.get("/invoices", response_model=List[Invoice])
async def get_invoices(
    current_user: dict = Depends(get_current_user),
    search: Optional[str] = None,
    type: Optional[str] = None,
    project_id: Optional[str] = None,
    status: Optional[str] = None,
    client_id: Optional[str] = None
):
    try:
        # Build query filters
        query_filter = {}
        
        if type:
            query_filter["invoice_type"] = type
        if project_id:
            query_filter["project_id"] = project_id
        if status:
            query_filter["status"] = status
        if client_id:
            query_filter["client_id"] = client_id
        
        invoices = await db.invoices.find(query_filter).to_list(1000)
        
        # Apply search filter if provided
        if search:
            search_lower = search.lower()
            filtered_invoices = []
            for invoice in invoices:
                if (search_lower in invoice.get("project_name", "").lower() or
                    search_lower in invoice.get("client_name", "").lower() or
                    search_lower in invoice.get("invoice_number", "").lower() or
                    search_lower in invoice.get("ra_number", "").lower()):
                    filtered_invoices.append(invoice)
            invoices = filtered_invoices
        
        # Filter and validate invoices to prevent validation errors
        valid_invoices = []
        for invoice in invoices:
            if not invoice or not isinstance(invoice, dict):
                continue
                
            try:
                # Ensure required fields exist with defaults
                cleaned_invoice = {
                    "id": invoice.get("id", str(uuid.uuid4())),
                    "invoice_number": invoice.get("invoice_number", ""),
                    "ra_number": invoice.get("ra_number", ""),
                    "project_id": invoice.get("project_id", ""),
                    "project_name": invoice.get("project_name", ""),
                    "client_id": invoice.get("client_id", ""),
                    "client_name": invoice.get("client_name", ""),
                    "invoice_type": invoice.get("invoice_type", "proforma"),
                    "items": [],
                    "subtotal": float(invoice.get("subtotal", 0)),
                    "total_gst_amount": float(invoice.get("total_gst_amount", invoice.get("gst_amount", 0))),
                    "total_amount": float(invoice.get("total_amount", 0)),
                    "is_partial": invoice.get("is_partial", True),
                    "billing_percentage": invoice.get("billing_percentage"),
                    "cumulative_billed": invoice.get("cumulative_billed"),
                    "status": invoice.get("status", "draft"),
                    "created_by": invoice.get("created_by"),
                    "reviewed_by": invoice.get("reviewed_by"),
                    "approved_by": invoice.get("approved_by"),
                    "invoice_date": invoice.get("invoice_date", datetime.utcnow()),
                    "due_date": invoice.get("due_date"),
                    "created_at": invoice.get("created_at", datetime.utcnow()),
                    "updated_at": invoice.get("updated_at", datetime.utcnow())
                }
                
                # Clean up items to match InvoiceItem model
                for item in invoice.get("items", []):
                    if isinstance(item, dict):
                        cleaned_item = {
                            "boq_item_id": item.get("boq_item_id", item.get("serial_number", str(uuid.uuid4()))),
                            "serial_number": str(item.get("serial_number", "")),
                            "description": str(item.get("description", "")),
                            "unit": str(item.get("unit", "nos")),
                            "quantity": float(item.get("quantity", 0)),
                            "rate": float(item.get("rate", 0)),
                            "amount": float(item.get("amount", 0)),
                            "gst_rate": float(item.get("gst_rate", 18.0)),
                            "gst_amount": float(item.get("gst_amount", 0)),
                            "total_with_gst": float(item.get("total_with_gst", item.get("amount", 0) * 1.18))
                        }
                        cleaned_invoice["items"].append(cleaned_item)
                
                valid_invoice = Invoice(**cleaned_invoice)
                valid_invoices.append(valid_invoice)
                
            except Exception as e:
                logger.warning(f"Skipping invalid invoice {invoice.get('id', 'unknown')}: {e}")
                continue
        
        return valid_invoices
        
    except Exception as e:
        logger.error(f"Error fetching invoices: {str(e)}")
        return []

@api_router.get("/invoices/{invoice_id}", response_model=Invoice)
async def get_invoice(invoice_id: str, current_user: dict = Depends(get_current_user)):
    """Get individual invoice by ID"""
    try:
        if not invoice_id or len(invoice_id.strip()) == 0:
            raise HTTPException(status_code=400, detail="Invoice ID is required")
        
        invoice_data = await db.invoices.find_one({"id": invoice_id})
        if not invoice_data:
            raise HTTPException(status_code=404, detail=f"Invoice with ID {invoice_id} not found")
        
        # Clean and validate invoice data
        cleaned_invoice = {
            "id": invoice_data.get("id", str(uuid.uuid4())),
            "invoice_number": invoice_data.get("invoice_number", ""),
            "ra_number": invoice_data.get("ra_number", ""),
            "project_id": invoice_data.get("project_id", ""),
            "project_name": invoice_data.get("project_name", ""),
            "client_id": invoice_data.get("client_id", ""),
            "client_name": invoice_data.get("client_name", ""),
            "invoice_type": invoice_data.get("invoice_type", "proforma"),
            "items": [],
            "subtotal": float(invoice_data.get("subtotal", 0)),
            "total_gst_amount": float(invoice_data.get("total_gst_amount", invoice_data.get("gst_amount", 0))),
            "total_amount": float(invoice_data.get("total_amount", 0)),
            "is_partial": invoice_data.get("is_partial", True),
            "billing_percentage": invoice_data.get("billing_percentage"),
            "cumulative_billed": invoice_data.get("cumulative_billed"),
            "status": invoice_data.get("status", "draft"),
            "created_by": invoice_data.get("created_by"),
            "reviewed_by": invoice_data.get("reviewed_by"),
            "approved_by": invoice_data.get("approved_by"),
            "invoice_date": invoice_data.get("invoice_date", datetime.utcnow()),
            "due_date": invoice_data.get("due_date"),
            "created_at": invoice_data.get("created_at", datetime.utcnow()),
            "updated_at": invoice_data.get("updated_at", datetime.utcnow())
        }
        
        # Clean up items
        for item in invoice_data.get("items", []):
            if isinstance(item, dict):
                cleaned_item = {
                    "boq_item_id": item.get("boq_item_id", item.get("serial_number", str(uuid.uuid4()))),
                    "serial_number": str(item.get("serial_number", "")),
                    "description": str(item.get("description", "")),
                    "unit": str(item.get("unit", "nos")),
                    "quantity": float(item.get("quantity", 0)),
                    "rate": float(item.get("rate", 0)),
                    "amount": float(item.get("amount", 0)),
                    "gst_rate": float(item.get("gst_rate", 18.0)),
                    "gst_amount": float(item.get("gst_amount", 0)),
                    "total_with_gst": float(item.get("total_with_gst", item.get("amount", 0) * 1.18))
                }
                cleaned_invoice["items"].append(cleaned_item)
        
        return Invoice(**cleaned_invoice)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching invoice {invoice_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@api_router.get("/invoices/{invoice_id}/pdf")
async def download_invoice_pdf(invoice_id: str, current_user: dict = Depends(get_current_user)):
    try:
        if not invoice_id or len(invoice_id.strip()) == 0:
            raise HTTPException(status_code=400, detail="Invoice ID is required")
        
        # Get invoice data
        invoice_data = await db.invoices.find_one({"id": invoice_id})
        if not invoice_data:
            raise HTTPException(status_code=404, detail=f"Invoice with ID {invoice_id} not found")
        
        # Get related project data
        project_data = await db.projects.find_one({"id": invoice_data.get("project_id")})
        if not project_data:
            # Create minimal project data if not found
            project_data = {
                "id": invoice_data.get("project_id", "unknown"),
                "project_name": invoice_data.get("project_name", "Unknown Project"),
                "architect": "Unknown Architect",
                "location": "Unknown Location",
                "client_id": invoice_data.get("client_id"),
                "boq_items": [],
                "total_project_value": 0,
                "advance_received": 0,
                "pending_payment": 0,
                "created_by": current_user["id"],
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
            
        # Get related client data
        client_data = await db.clients.find_one({"id": invoice_data.get("client_id")})
        if not client_data:
            # Create minimal client data if not found
            client_data = {
                "id": invoice_data.get("client_id", "unknown"),
                "name": invoice_data.get("client_name", "Unknown Client"),
                "bill_to_address": "Unknown Address",
                "ship_to_address": "Unknown Address",
                "gst_no": "",
                "contact_person": "",
                "phone": "",
                "email": "",
                "created_at": datetime.utcnow()
            }
        
        # Clean and validate data before PDF generation
        try:
            # Ensure required fields exist with defaults
            cleaned_invoice = {
                "id": invoice_data.get("id", invoice_id),
                "invoice_number": invoice_data.get("invoice_number", "Unknown"),
                "ra_number": invoice_data.get("ra_number", "Unknown"),
                "project_id": invoice_data.get("project_id", ""),
                "project_name": invoice_data.get("project_name", "Unknown Project"),
                "client_id": invoice_data.get("client_id", ""),
                "client_name": invoice_data.get("client_name", "Unknown Client"),
                "invoice_type": invoice_data.get("invoice_type", "proforma"),
                "items": [],
                "subtotal": float(invoice_data.get("subtotal", 0)),
                "total_gst_amount": float(invoice_data.get("total_gst_amount", 0)),
                "total_amount": float(invoice_data.get("total_amount", 0)),
                "is_partial": invoice_data.get("is_partial", True),
                "billing_percentage": invoice_data.get("billing_percentage"),
                "cumulative_billed": invoice_data.get("cumulative_billed"),
                "status": invoice_data.get("status", "draft"),
                "created_by": invoice_data.get("created_by"),
                "reviewed_by": invoice_data.get("reviewed_by"),
                "approved_by": invoice_data.get("approved_by"),
                "invoice_date": invoice_data.get("invoice_date", datetime.utcnow()),
                "due_date": invoice_data.get("due_date"),
                "created_at": invoice_data.get("created_at", datetime.utcnow()),
                "updated_at": invoice_data.get("updated_at", datetime.utcnow())
            }
            
            # Clean items data
            for item in invoice_data.get("items", []):
                if isinstance(item, dict):
                    cleaned_item = {
                        "boq_item_id": item.get("boq_item_id", item.get("serial_number", str(len(cleaned_invoice["items"]) + 1))),
                        "serial_number": str(item.get("serial_number", len(cleaned_invoice["items"]) + 1)),
                        "description": str(item.get("description", "Unknown Item")),
                        "unit": str(item.get("unit", "nos")),
                        "quantity": float(item.get("quantity", 0)),
                        "rate": float(item.get("rate", 0)),
                        "amount": float(item.get("amount", 0)),
                        "gst_rate": float(item.get("gst_rate", 18.0)),
                        "gst_amount": float(item.get("gst_amount", 0)),
                        "total_with_gst": float(item.get("total_with_gst", 0))
                    }
                    cleaned_invoice["items"].append(cleaned_item)
            
            # Create model instances with error handling
            invoice = Invoice(**cleaned_invoice)
            project = Project(**project_data)
            client = ClientInfo(**client_data)
            
        except Exception as validation_error:
            logger.error(f"Data validation error for invoice {invoice_id}: {str(validation_error)}")
            # Return a simple error PDF instead of failing
            error_pdf = create_error_pdf(f"Invoice {invoice_data.get('invoice_number', 'Unknown')}", str(validation_error))
            return Response(
                content=error_pdf,
                media_type="application/pdf",
                headers={
                    "Content-Type": "application/pdf",
                    "Content-Disposition": f"inline; filename=invoice_error_{invoice_id}.pdf",
                    "Cache-Control": "no-cache"
                }
            )
        
        # Generate PDF
        try:
            pdf_generator = PDFGenerator()
            pdf_buffer = await pdf_generator.generate_invoice_pdf(invoice, project, client)
            
            # Convert BytesIO to bytes for response
            pdf_content = pdf_buffer.getvalue()
            
            if len(pdf_content) < 100:  # Check if PDF is too small (likely an error)
                raise Exception("Generated PDF is too small - likely incomplete")
            
            await log_activity(
                current_user["id"], current_user["email"], current_user["role"],
                "invoice_downloaded", f"Downloaded PDF for invoice: {invoice.invoice_number}",
                invoice_id=invoice_id
            )
            
            # Return as inline response for viewing
            return Response(
                content=pdf_content,
                media_type="application/pdf",
                headers={
                    "Content-Type": "application/pdf",
                    "Content-Disposition": f"inline; filename=invoice_{invoice.invoice_number}.pdf",
                    "Content-Length": str(len(pdf_content)),
                    "Cache-Control": "no-cache"
                }
            )
            
        except Exception as pdf_error:
            logger.error(f"PDF generation error for invoice {invoice_id}: {str(pdf_error)}")
            # Return a simple error PDF
            error_pdf = create_error_pdf(f"Invoice {invoice_data.get('invoice_number', 'Unknown')}", f"PDF generation failed: {str(pdf_error)}")
            return Response(
                content=error_pdf,
                media_type="application/pdf",
                headers={
                    "Content-Type": "application/pdf",
                    "Content-Disposition": f"inline; filename=invoice_error_{invoice_id}.pdf",
                    "Cache-Control": "no-cache"
                }
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error downloading invoice PDF {invoice_id}: {str(e)}")
        # Return error PDF instead of HTTP error
        try:
            error_pdf = create_error_pdf("Invoice Error", f"Failed to generate PDF: {str(e)}")
            return Response(
                content=error_pdf,
                media_type="application/pdf",
                headers={
                    "Content-Type": "application/pdf",
                    "Content-Disposition": f"inline; filename=invoice_error_{invoice_id}.pdf",
                    "Cache-Control": "no-cache"
                }
            )
        except:
            raise HTTPException(status_code=500, detail=f"Critical error: {str(e)}")

def create_error_pdf(title: str, error_message: str) -> bytes:
    """Create a simple error PDF when regular PDF generation fails"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        import io
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                              rightMargin=inch, leftMargin=inch, 
                              topMargin=inch, bottomMargin=inch)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.red,
            alignment=1  # Center alignment
        )
        
        error_style = ParagraphStyle(
            'ErrorStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=20,
            textColor=colors.black
        )
        
        content = []
        content.append(Paragraph(f"Error: {title}", title_style))
        content.append(Spacer(1, 20))
        content.append(Paragraph(f"Details: {error_message}", error_style))
        content.append(Spacer(1, 20))
        content.append(Paragraph("Please contact support for assistance.", error_style))
        
        doc.build(content)
        buffer.seek(0)
        return buffer.read()
        
    except Exception as e:
        # If even error PDF creation fails, return minimal PDF
        return b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/Resources <<\n/Font <<\n/F1 4 0 R\n>>\n>>\n/MediaBox [0 0 612 792]\n/Contents 5 0 R\n>>\nendobj\n4 0 obj\n<<\n/Type /Font\n/Subtype /Type1\n/BaseFont /Times-Roman\n>>\nendobj\n5 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(PDF Generation Error) Tj\nET\nendstream\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000079 00000 n \n0000000173 00000 n \n0000000301 00000 n \n0000000380 00000 n \ntrailer\n<<\n/Size 6\n/Root 1 0 R\n>>\nstartxref\n492\n%%EOF"

@api_router.get("/dashboard/stats")
async def get_dashboard_stats(current_user: dict = Depends(get_current_user)):
    total_projects = await db.projects.count_documents({})
    total_invoices = await db.invoices.count_documents({})
    
    # Calculate financial stats
    pipeline = [
        {"$group": {
            "_id": None,
            "total_invoiced": {"$sum": "$total_amount"},
            "total_advance": {"$sum": "$advance_received"}
        }}
    ]
    
    financial_stats = await db.projects.aggregate(pipeline).to_list(1)
    total_invoiced = financial_stats[0]["total_invoiced"] if financial_stats else 0
    total_advance = financial_stats[0]["total_advance"] if financial_stats else 0
    
    pending_payment = total_invoiced - total_advance
    
    return {
        "total_projects": total_projects,
        "total_invoices": total_invoices,
        "total_invoiced_value": total_invoiced,
        "advance_received": total_advance,
        "pending_payment": pending_payment
    }

@api_router.get("/activity-logs")
async def get_activity_logs(
    skip: int = 0, 
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    if current_user["role"] != UserRole.SUPER_ADMIN:
        raise HTTPException(status_code=403, detail="Only super admin can view logs")
    
    logs = await db.activity_logs.find().skip(skip).limit(limit).sort("timestamp", -1).to_list(limit)
    return [ActivityLog(**log) for log in logs]

# Item Master Management
@api_router.post("/item-master", response_model=dict)
async def create_master_item(item_data: MasterItem, current_user: dict = Depends(get_current_user)):
    """Create a new master item"""
    try:
        item_data.created_by = current_user["id"]
        item_data.updated_at = datetime.utcnow()
        
        # Check if similar item already exists
        import re as regex_module
        escaped_description = regex_module.escape(item_data.description)
        existing_item = await db.master_items.find_one({
            "description": {"$regex": f"^{escaped_description}$", "$options": "i"},
            "unit": item_data.unit
        })
        
        if existing_item:
            raise HTTPException(status_code=400, detail="Similar item already exists in master")
        
        await db.master_items.insert_one(item_data.dict())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "master_item_created", f"Created master item: {item_data.description} ({item_data.unit})"
        )
        
        return {"message": "Master item created successfully", "item_id": item_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating master item: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create master item: {str(e)}")

@api_router.get("/item-master", response_model=List[MasterItem])
async def get_master_items(
    category: Optional[str] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get all master items with optional filtering"""
    try:
        query = {}
        
        if category:
            query["category"] = category
            
        if search:
            query["$or"] = [
                {"description": {"$regex": search, "$options": "i"}},
                {"category": {"$regex": search, "$options": "i"}}
            ]
        
        items = await db.master_items.find(query).sort("description", 1).to_list(1000)
        return [MasterItem(**item) for item in items]
        
    except Exception as e:
        logger.error(f"Error fetching master items: {str(e)}")
        return []

@api_router.put("/item-master/{item_id}", response_model=dict)
async def update_master_item(
    item_id: str, 
    updated_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update a master item"""
    try:
        if not item_id:
            raise HTTPException(status_code=400, detail="Item ID is required")
        
        item = await db.master_items.find_one({"id": item_id})
        if not item:
            raise HTTPException(status_code=404, detail="Master item not found")
        
        # Update only allowed fields
        allowed_fields = ["description", "unit", "standard_rate", "category"]
        update_data = {k: v for k, v in updated_data.items() if k in allowed_fields}
        update_data["updated_at"] = datetime.utcnow()
        
        await db.master_items.update_one(
            {"id": item_id},
            {"$set": update_data}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "master_item_updated", f"Updated master item: {item['description']}"
        )
        
        return {"message": "Master item updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating master item: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update master item: {str(e)}")

@api_router.delete("/item-master/{item_id}")
async def delete_master_item(item_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a master item"""
    try:
        if current_user["role"] not in [UserRole.SUPER_ADMIN, UserRole.INVOICE_CREATOR]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        
        item = await db.master_items.find_one({"id": item_id})
        if not item:
            raise HTTPException(status_code=404, detail="Master item not found")
        
        await db.master_items.delete_one({"id": item_id})
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "master_item_deleted", f"Deleted master item: {item['description']}"
        )
        
        return {"message": "Master item deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting master item: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete master item: {str(e)}")

@api_router.post("/item-master/auto-populate")
async def auto_populate_master_items(current_user: dict = Depends(get_current_user)):
    """Auto-populate master items from existing BOQ items across all projects"""
    try:
        if current_user["role"] not in [UserRole.SUPER_ADMIN, UserRole.INVOICE_CREATOR]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        
        # Get all projects and extract unique BOQ items
        projects = await db.projects.find().to_list(1000)
        unique_items = {}
        
        for project in projects:
            for boq_item in project.get("boq_items", []):
                description = boq_item.get("description", "").strip()
                unit = boq_item.get("unit", "nos").strip()
                rate = boq_item.get("rate", 0)
                
                if description and len(description) > 3:
                    key = f"{description.lower()}_{unit.lower()}"
                    
                    if key not in unique_items:
                        unique_items[key] = {
                            "description": description,
                            "unit": unit,
                            "rates": [rate],
                            "count": 1
                        }
                    else:
                        unique_items[key]["rates"].append(rate)
                        unique_items[key]["count"] += 1
        
        # Create master items
        created_count = 0
        for item_data in unique_items.values():
            # Calculate average rate
            valid_rates = [r for r in item_data["rates"] if r > 0]
            avg_rate = sum(valid_rates) / len(valid_rates) if valid_rates else 0
            
            # Check if item already exists
            import re as regex_module
            escaped_description = regex_module.escape(item_data['description'])
            existing = await db.master_items.find_one({
                "description": {"$regex": f"^{escaped_description}$", "$options": "i"},
                "unit": item_data["unit"]
            })
            
            if not existing:
                master_item = MasterItem(
                    description=item_data["description"],
                    unit=item_data["unit"],
                    standard_rate=avg_rate,
                    usage_count=item_data["count"],
                    created_by=current_user["id"]
                )
                
                await db.master_items.insert_one(master_item.dict())
                created_count += 1
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "master_items_auto_populated", f"Auto-populated {created_count} master items from existing BOQ data"
        )
        
        return {
            "message": f"Successfully created {created_count} master items",
            "created_count": created_count,
            "total_unique_items": len(unique_items)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error auto-populating master items: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to auto-populate master items: {str(e)}")

# Smart Search and Filters
@api_router.get("/search")
async def global_search(
    query: str,
    entity_type: Optional[str] = None,  # projects, clients, invoices, all
    limit: int = 20,
    current_user: dict = Depends(get_current_user)
):
    """Global search across projects, clients, and invoices"""
    try:
        results = {"projects": [], "clients": [], "invoices": [], "total_count": 0}
        
        search_regex = {"$regex": query, "$options": "i"}
        
        # Search projects
        if not entity_type or entity_type in ["projects", "all"]:
            project_query = {"$or": [
                {"project_name": search_regex},
                {"client_name": search_regex},
                {"architect": search_regex}
            ]}
            projects = await db.projects.find(project_query).limit(limit).to_list(limit)
            results["projects"] = [
                {
                    "id": p.get("id"),
                    "project_name": p.get("project_name"),
                    "client_name": p.get("client_name"),
                    "architect": p.get("architect"),
                    "total_value": p.get("total_project_value", 0),
                    "type": "project"
                } for p in projects
            ]
        
        # Search clients
        if not entity_type or entity_type in ["clients", "all"]:
            client_query = {"$or": [
                {"name": search_regex},
                {"bill_to_address": search_regex},
                {"gst_no": search_regex}
            ]}
            clients = await db.clients.find(client_query).limit(limit).to_list(limit)
            results["clients"] = [
                {
                    "id": c.get("id"),
                    "name": c.get("name"),
                    "bill_to_address": c.get("bill_to_address"),
                    "gst_no": c.get("gst_no"),
                    "type": "client"
                } for c in clients
            ]
        
        # Search invoices
        if not entity_type or entity_type in ["invoices", "all"]:
            invoice_query = {"$or": [
                {"invoice_number": search_regex},
                {"ra_number": search_regex},
                {"project_name": search_regex},
                {"client_name": search_regex}
            ]}
            invoices = await db.invoices.find(invoice_query).limit(limit).to_list(limit)
            results["invoices"] = [
                {
                    "id": i.get("id"),
                    "invoice_number": i.get("invoice_number"),
                    "ra_number": i.get("ra_number"),
                    "project_name": i.get("project_name"),
                    "client_name": i.get("client_name"),
                    "total_amount": i.get("total_amount", 0),
                    "status": i.get("status"),
                    "type": "invoice"
                } for i in invoices
            ]
        
        results["total_count"] = len(results["projects"]) + len(results["clients"]) + len(results["invoices"])
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "global_search", f"Performed global search for: {query}"
        )
        
        return results
        
    except Exception as e:
        logger.error(f"Error in global search: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

@api_router.get("/filters/projects")
async def get_filtered_projects(
    client_id: Optional[str] = None,
    architect: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    min_value: Optional[float] = None,
    max_value: Optional[float] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get filtered projects with advanced filtering"""
    try:
        query = {}
        
        if client_id:
            query["client_id"] = client_id
            
        if architect:
            query["architect"] = {"$regex": architect, "$options": "i"}
            
        if date_from or date_to:
            date_query = {}
            if date_from:
                date_query["$gte"] = datetime.fromisoformat(date_from.replace('Z', '+00:00'))
            if date_to:
                date_query["$lte"] = datetime.fromisoformat(date_to.replace('Z', '+00:00'))
            query["created_at"] = date_query
            
        if min_value or max_value:
            value_query = {}
            if min_value:
                value_query["$gte"] = min_value
            if max_value:
                value_query["$lte"] = max_value
            query["total_project_value"] = value_query
            
        projects = await db.projects.find(query).sort("created_at", -1).to_list(1000)
        
        # Filter and validate projects (same as get_projects endpoint)
        valid_projects = []
        for project in projects:
            if not project or not isinstance(project, dict):
                continue
                
            cleaned_project = {
                "id": project.get("id", str(uuid.uuid4())),
                "project_name": project.get("project_name", "Untitled Project"),
                "architect": project.get("architect", "Unknown Architect"),
                "client_id": project.get("client_id", ""),
                "client_name": project.get("client_name", "Unknown Client"),
                "metadata": project.get("metadata", {}),
                "boq_items": project.get("boq_items", []),
                "total_project_value": float(project.get("total_project_value", 0)),
                "advance_received": float(project.get("advance_received", 0)),
                "pending_payment": float(project.get("pending_payment", 0)),
                "created_by": project.get("created_by"),
                "created_at": project.get("created_at", datetime.utcnow()),
                "updated_at": project.get("updated_at", datetime.utcnow())
            }
            
            try:
                valid_project = Project(**cleaned_project)
                valid_projects.append(valid_project)
            except Exception as e:
                logger.warning(f"Skipping invalid project {project.get('id', 'unknown')}: {e}")
                continue
        
        return valid_projects
        
    except Exception as e:
        logger.error(f"Error filtering projects: {str(e)}")
        return []

@api_router.get("/filters/invoices")
async def get_filtered_invoices(
    status: Optional[str] = None,
    invoice_type: Optional[str] = None,
    project_id: Optional[str] = None,
    client_id: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    min_amount: Optional[float] = None,
    max_amount: Optional[float] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get filtered invoices with advanced filtering"""
    try:
        query = {}
        
        if status:
            query["status"] = status
            
        if invoice_type:
            query["invoice_type"] = invoice_type
            
        if project_id:
            query["project_id"] = project_id
            
        if client_id:
            query["client_id"] = client_id
            
        if date_from or date_to:
            date_query = {}
            if date_from:
                date_query["$gte"] = datetime.fromisoformat(date_from.replace('Z', '+00:00'))
            if date_to:
                date_query["$lte"] = datetime.fromisoformat(date_to.replace('Z', '+00:00'))
            query["created_at"] = date_query
            
        if min_amount or max_amount:
            amount_query = {}
            if min_amount:
                amount_query["$gte"] = min_amount
            if max_amount:
                amount_query["$lte"] = max_amount
            query["total_amount"] = amount_query
        
        invoices = await db.invoices.find(query).sort("created_at", -1).to_list(1000)
        
        # Use the same validation logic as get_invoices endpoint
        valid_invoices = []
        for invoice in invoices:
            if not invoice or not isinstance(invoice, dict):
                continue
                
            try:
                cleaned_invoice = {
                    "id": invoice.get("id", str(uuid.uuid4())),
                    "invoice_number": invoice.get("invoice_number", ""),
                    "ra_number": invoice.get("ra_number", ""),
                    "project_id": invoice.get("project_id", ""),
                    "project_name": invoice.get("project_name", ""),
                    "client_id": invoice.get("client_id", ""),
                    "client_name": invoice.get("client_name", ""),
                    "invoice_type": invoice.get("invoice_type", "proforma"),
                    "items": [],
                    "subtotal": float(invoice.get("subtotal", 0)),
                    "total_gst_amount": float(invoice.get("total_gst_amount", invoice.get("gst_amount", 0))),
                    "total_amount": float(invoice.get("total_amount", 0)),
                    "is_partial": invoice.get("is_partial", True),
                    "billing_percentage": invoice.get("billing_percentage"),
                    "cumulative_billed": invoice.get("cumulative_billed"),
                    "status": invoice.get("status", "draft"),
                    "created_by": invoice.get("created_by"),
                    "reviewed_by": invoice.get("reviewed_by"),
                    "approved_by": invoice.get("approved_by"),
                    "invoice_date": invoice.get("invoice_date", datetime.utcnow()),
                    "due_date": invoice.get("due_date"),
                    "created_at": invoice.get("created_at", datetime.utcnow()),
                    "updated_at": invoice.get("updated_at", datetime.utcnow())
                }
                
                # Clean up items
                for item in invoice.get("items", []):
                    if isinstance(item, dict):
                        cleaned_item = {
                            "boq_item_id": item.get("boq_item_id", item.get("serial_number", str(uuid.uuid4()))),
                            "serial_number": str(item.get("serial_number", "")),
                            "description": str(item.get("description", "")),
                            "unit": str(item.get("unit", "nos")),
                            "quantity": float(item.get("quantity", 0)),
                            "rate": float(item.get("rate", 0)),
                            "amount": float(item.get("amount", 0)),
                            "gst_rate": float(item.get("gst_rate", 18.0)),
                            "gst_amount": float(item.get("gst_amount", 0)),
                            "total_with_gst": float(item.get("total_with_gst", item.get("amount", 0) * 1.18))
                        }
                        cleaned_invoice["items"].append(cleaned_item)
                
                valid_invoice = Invoice(**cleaned_invoice)
                valid_invoices.append(valid_invoice)
                
            except Exception as e:
                logger.warning(f"Skipping invalid invoice {invoice.get('id', 'unknown')}: {e}")
                continue
        
        return valid_invoices
        
    except Exception as e:
        logger.error(f"Error filtering invoices: {str(e)}")
        return []

# Reports and Insights
@api_router.get("/reports/gst-summary")
async def get_gst_summary(
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get GST summary report with tax breakdown"""
    try:
        query = {}
        if date_from or date_to:
            date_query = {}
            if date_from:
                date_query["$gte"] = datetime.fromisoformat(date_from.replace('Z', '+00:00'))
            if date_to:
                date_query["$lte"] = datetime.fromisoformat(date_to.replace('Z', '+00:00'))
            query["created_at"] = date_query
        
        invoices = await db.invoices.find(query).to_list(1000)
        
        gst_summary = {
            "total_invoices": len(invoices),
            "total_taxable_amount": 0,
            "total_gst_amount": 0,
            "total_amount_with_gst": 0,
            "gst_breakdown": {},
            "monthly_breakdown": {},
            "invoice_type_breakdown": {"proforma": 0, "tax_invoice": 0}
        }
        
        for invoice in invoices:
            subtotal = invoice.get("subtotal", 0)
            gst_amount = invoice.get("total_gst_amount", invoice.get("gst_amount", 0))
            total_amount = invoice.get("total_amount", 0)
            invoice_type = invoice.get("invoice_type", "proforma")
            
            gst_summary["total_taxable_amount"] += subtotal
            gst_summary["total_gst_amount"] += gst_amount
            gst_summary["total_amount_with_gst"] += total_amount
            gst_summary["invoice_type_breakdown"][invoice_type] += total_amount
            
            # Monthly breakdown
            created_at = invoice.get("created_at", datetime.utcnow())
            if isinstance(created_at, str):
                created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            
            month_key = created_at.strftime("%Y-%m")
            if month_key not in gst_summary["monthly_breakdown"]:
                gst_summary["monthly_breakdown"][month_key] = {
                    "month": created_at.strftime("%B %Y"),
                    "total_invoices": 0,
                    "taxable_amount": 0,
                    "gst_amount": 0,
                    "total_amount": 0
                }
            
            gst_summary["monthly_breakdown"][month_key]["total_invoices"] += 1
            gst_summary["monthly_breakdown"][month_key]["taxable_amount"] += subtotal
            gst_summary["monthly_breakdown"][month_key]["gst_amount"] += gst_amount
            gst_summary["monthly_breakdown"][month_key]["total_amount"] += total_amount
            
            # GST rate breakdown
            for item in invoice.get("items", []):
                gst_rate = item.get("gst_rate", 18.0)
                item_gst_amount = item.get("gst_amount", 0)
                
                if gst_rate not in gst_summary["gst_breakdown"]:
                    gst_summary["gst_breakdown"][gst_rate] = {
                        "rate": gst_rate,
                        "taxable_amount": 0,
                        "gst_amount": 0,
                        "total_amount": 0
                    }
                
                gst_summary["gst_breakdown"][gst_rate]["taxable_amount"] += item.get("amount", 0)
                gst_summary["gst_breakdown"][gst_rate]["gst_amount"] += item_gst_amount
                gst_summary["gst_breakdown"][gst_rate]["total_amount"] += item.get("total_with_gst", 0)
        
        # Convert dict to list for better frontend handling
        gst_summary["monthly_breakdown"] = list(gst_summary["monthly_breakdown"].values())
        gst_summary["gst_breakdown"] = list(gst_summary["gst_breakdown"].values())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "gst_report_generated", f"Generated GST summary report"
        )
        
        return gst_summary
        
    except Exception as e:
        logger.error(f"Error generating GST summary: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate GST summary: {str(e)}")

@api_router.get("/reports/insights")
async def get_insights(current_user: dict = Depends(get_current_user)):
    """Get business insights and analytics"""
    try:
        # Get current date
        now = datetime.utcnow()
        current_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        last_month = (current_month - timedelta(days=1)).replace(day=1)
        
        # Basic counts
        total_projects = await db.projects.count_documents({})
        total_clients = await db.clients.count_documents({})
        total_invoices = await db.invoices.count_documents({})
        
        # Financial data
        projects = await db.projects.find().to_list(1000)
        invoices = await db.invoices.find().to_list(1000)
        
        # Calculate project insights
        total_project_value = sum(p.get("total_project_value", 0) for p in projects)
        total_advance_received = sum(p.get("advance_received", 0) for p in projects)
        total_pending_payment = sum(p.get("pending_payment", 0) for p in projects)
        
        # Calculate invoice insights
        total_invoiced_value = sum(i.get("total_amount", 0) for i in invoices)
        
        # Monthly trends (last 6 months)
        monthly_data = {}
        for i in range(6):
            month_start = (current_month - timedelta(days=i * 30)).replace(day=1)
            month_end = (month_start + timedelta(days=31)).replace(day=1) - timedelta(seconds=1)
            month_key = month_start.strftime("%Y-%m")
            
            monthly_invoices = [
                inv for inv in invoices 
                if month_start <= datetime.fromisoformat(str(inv.get("created_at", now)).replace('Z', '+00:00')) <= month_end
            ]
            
            monthly_data[month_key] = {
                "month": month_start.strftime("%B %Y"),
                "invoices_count": len(monthly_invoices),
                "invoices_value": sum(inv.get("total_amount", 0) for inv in monthly_invoices)
            }
        
        # Top clients by value
        client_values = {}
        for project in projects:
            client_name = project.get("client_name", "Unknown")
            client_values[client_name] = client_values.get(client_name, 0) + project.get("total_project_value", 0)
        
        top_clients = sorted(
            [{"name": k, "total_value": v} for k, v in client_values.items()],
            key=lambda x: x["total_value"],
            reverse=True
        )[:5]
        
        # Active users (users who performed actions in last 30 days)
        thirty_days_ago = now - timedelta(days=30)
        recent_logs = await db.activity_logs.find({
            "timestamp": {"$gte": thirty_days_ago}
        }).to_list(1000)
        
        active_users = len(set(log.get("user_email") for log in recent_logs))
        
        insights = {
            "overview": {
                "total_projects": total_projects,
                "total_clients": total_clients,
                "total_invoices": total_invoices,
                "active_users": active_users
            },
            "financial": {
                "total_project_value": total_project_value,
                "total_advance_received": total_advance_received,
                "total_pending_payment": total_pending_payment,
                "total_invoiced_value": total_invoiced_value,
                "collection_percentage": (total_advance_received / total_project_value * 100) if total_project_value > 0 else 0
            },
            "trends": {
                "monthly_data": list(monthly_data.values()),
                "top_clients": top_clients
            },
            "performance": {
                "avg_project_value": total_project_value / total_projects if total_projects > 0 else 0,
                "avg_invoice_value": total_invoiced_value / total_invoices if total_invoices > 0 else 0,
                "projects_per_client": total_projects / total_clients if total_clients > 0 else 0
            }
        }
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "insights_report_generated", f"Generated business insights report"
        )
        
        return insights
        
    except Exception as e:
        logger.error(f"Error generating insights: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate insights: {str(e)}")

@api_router.get("/reports/client-summary/{client_id}")
async def get_client_summary(client_id: str, current_user: dict = Depends(get_current_user)):
    """Get detailed summary for a specific client"""
    try:
        client = await db.clients.find_one({"id": client_id})
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
        
        # Get client's projects and invoices
        projects = await db.projects.find({"client_id": client_id}).to_list(1000)
        invoices = await db.invoices.find({"client_id": client_id}).to_list(1000)
        
        # Clean up data to remove ObjectIds and ensure JSON serialization
        def clean_data(data):
            if isinstance(data, list):
                return [clean_data(item) for item in data]
            elif isinstance(data, dict):
                cleaned = {}
                for k, v in data.items():
                    if k == '_id':  # Skip MongoDB ObjectId
                        continue
                    cleaned[k] = clean_data(v)
                return cleaned
            else:
                return data
        
        clean_client = clean_data(client)
        clean_projects = clean_data(projects)
        clean_invoices = clean_data(invoices)
        
        summary = {
            "client_info": clean_client,
            "projects_count": len(clean_projects),
            "invoices_count": len(clean_invoices),
            "total_project_value": sum(p.get("total_project_value", 0) for p in clean_projects),
            "total_invoiced_value": sum(i.get("total_amount", 0) for i in clean_invoices),
            "total_advance_received": sum(p.get("advance_received", 0) for p in clean_projects),
            "pending_amount": sum(p.get("pending_payment", 0) for p in clean_projects),
            "projects": clean_projects,
            "recent_invoices": sorted(clean_invoices, key=lambda x: x.get("created_at", datetime.min), reverse=True)[:5]
        }
        
        return summary
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating client summary: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate client summary: {str(e)}")

# Bank Guarantee Management
@api_router.post("/bank-guarantees", response_model=dict)
async def create_bank_guarantee(guarantee_data: BankGuarantee, current_user: dict = Depends(get_current_user)):
    """Create a new bank guarantee"""
    try:
        # Validate project exists
        project = await db.projects.find_one({"id": guarantee_data.project_id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        guarantee_data.created_by = current_user["id"]
        guarantee_data.updated_at = datetime.utcnow()
        
        await db.bank_guarantees.insert_one(guarantee_data.dict())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "bank_guarantee_created", 
            f"Created {guarantee_data.guarantee_type} bank guarantee for {guarantee_data.project_name} - ₹{guarantee_data.guarantee_amount:,.2f}"
        )
        
        return {"message": "Bank guarantee created successfully", "guarantee_id": guarantee_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating bank guarantee: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create bank guarantee: {str(e)}")

@api_router.get("/bank-guarantees", response_model=List[dict])
async def get_bank_guarantees(
    project_id: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get bank guarantees with optional filtering"""
    try:
        query = {}
        
        if project_id:
            query["project_id"] = project_id
            
        if status:
            query["status"] = status
        
        guarantees = await db.bank_guarantees.find(query).sort("created_at", -1).to_list(1000)
        
        # Add expiry status
        current_date = datetime.utcnow()
        for guarantee in guarantees:
            validity_date = guarantee.get("validity_date")
            if isinstance(validity_date, str):
                validity_date = datetime.fromisoformat(validity_date.replace('Z', '+00:00'))
            
            if validity_date and validity_date < current_date and guarantee.get("status") == "active":
                # Update status to expired
                await db.bank_guarantees.update_one(
                    {"id": guarantee["id"]},
                    {"$set": {"status": "expired", "updated_at": datetime.utcnow()}}
                )
                guarantee["status"] = "expired"
        
        return guarantees
        
    except Exception as e:
        logger.error(f"Error fetching bank guarantees: {str(e)}")
        return []

@api_router.put("/bank-guarantees/{guarantee_id}", response_model=dict)
async def update_bank_guarantee(
    guarantee_id: str,
    update_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update bank guarantee"""
    try:
        guarantee = await db.bank_guarantees.find_one({"id": guarantee_id})
        if not guarantee:
            raise HTTPException(status_code=404, detail="Bank guarantee not found")
        
        # Update allowed fields
        allowed_fields = ["guarantee_amount", "guarantee_percentage", "validity_date", "status", "guarantee_details"]
        update_fields = {k: v for k, v in update_data.items() if k in allowed_fields}
        update_fields["updated_at"] = datetime.utcnow()
        
        await db.bank_guarantees.update_one(
            {"id": guarantee_id},
            {"$set": update_fields}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "bank_guarantee_updated", 
            f"Updated bank guarantee {guarantee['guarantee_number']} for {guarantee['project_name']}"
        )
        
        return {"message": "Bank guarantee updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating bank guarantee: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update bank guarantee: {str(e)}")

@api_router.get("/bank-guarantees/summary")
async def get_bank_guarantees_summary(current_user: dict = Depends(get_current_user)):
    """Get bank guarantees summary"""
    try:
        guarantees = await db.bank_guarantees.find().to_list(1000)
        
        summary = {
            "total_guarantees": len(guarantees),
            "active_guarantees": len([g for g in guarantees if g.get("status") == "active"]),
            "expired_guarantees": len([g for g in guarantees if g.get("status") == "expired"]),
            "total_guarantee_amount": sum(g.get("guarantee_amount", 0) for g in guarantees if g.get("status") == "active"),
            "guarantees_by_type": {},
            "expiring_soon": []
        }
        
        # Group by guarantee type
        for guarantee in guarantees:
            gtype = guarantee.get("guarantee_type", "Unknown")
            if gtype not in summary["guarantees_by_type"]:
                summary["guarantees_by_type"][gtype] = {
                    "count": 0,
                    "total_amount": 0
                }
            summary["guarantees_by_type"][gtype]["count"] += 1
            summary["guarantees_by_type"][gtype]["total_amount"] += guarantee.get("guarantee_amount", 0)
        
        # Find guarantees expiring in next 30 days
        current_date = datetime.utcnow()
        thirty_days_later = current_date + timedelta(days=30)
        
        for guarantee in guarantees:
            if guarantee.get("status") != "active":
                continue
                
            validity_date = guarantee.get("validity_date")
            if isinstance(validity_date, str):
                validity_date = datetime.fromisoformat(validity_date.replace('Z', '+00:00'))
            
            if validity_date and current_date < validity_date < thirty_days_later:
                summary["expiring_soon"].append({
                    "id": guarantee["id"],
                    "project_name": guarantee["project_name"],
                    "guarantee_number": guarantee["guarantee_number"],
                    "guarantee_amount": guarantee["guarantee_amount"],
                    "validity_date": validity_date,
                    "days_remaining": (validity_date - current_date).days
                })
        
        return summary
        
    except Exception as e:
        logger.error(f"Error generating bank guarantees summary: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate summary: {str(e)}")

# PDF Text Extraction Engine for PO Processing
@api_router.post("/pdf-processor/extract")
async def extract_pdf_data(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Extract data from uploaded PDF/DOCX Purchase Order"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Read file content
        file_content = await file.read()
        
        if len(file_content) == 0:
            raise HTTPException(status_code=400, detail="Empty file provided")
        
        # Initialize PDF parser
        parser = POPDFParser()
        
        # Extract data
        extracted_data = await parser.extract_from_file(file_content, file.filename)
        
    except ValueError as e:
        # Handle unsupported file formats and other value errors
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing PDF file {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

    try:
        extraction_record = {
            "id": str(uuid.uuid4()),
            "original_filename": file.filename,
            "extracted_data": extracted_data.dict(),
            "processed_by": current_user["id"],
            "processed_at": datetime.utcnow(),
            "file_size": len(file_content)
        }
        
        await db.pdf_extractions.insert_one(extraction_record)
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "pdf_data_extracted", 
            f"Extracted data from PO file: {file.filename} (Method: {extracted_data.extraction_method}, Confidence: {extracted_data.confidence_score:.2f})"
        )
        
        return {
            "extraction_id": extraction_record["id"],
            "extracted_data": extracted_data,
            "processing_info": {
                "filename": file.filename,
                "file_size": len(file_content),
                "extraction_method": extracted_data.extraction_method,
                "confidence_score": extracted_data.confidence_score
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing PDF file {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

@api_router.get("/pdf-processor/extractions")
async def get_pdf_extractions(
    skip: int = 0,
    limit: int = 50,
    current_user: dict = Depends(get_current_user)
):
    """Get list of PDF extractions"""
    try:
        extractions = await db.pdf_extractions.find().skip(skip).limit(limit).sort("processed_at", -1).to_list(limit)
        
        # Convert ObjectId to string for JSON serialization
        for extraction in extractions:
            if "_id" in extraction:
                extraction["_id"] = str(extraction["_id"])
        
        return {
            "extractions": extractions,
            "total": len(extractions)
        }
        
    except Exception as e:
        logger.error(f"Error fetching PDF extractions: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch extractions: {str(e)}")

@api_router.get("/pdf-processor/extractions/{extraction_id}")
async def get_pdf_extraction(
    extraction_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get specific PDF extraction details"""
    try:
        extraction = await db.pdf_extractions.find_one({"id": extraction_id})
        
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")
        
        # Convert ObjectId to string for JSON serialization
        if "_id" in extraction:
            extraction["_id"] = str(extraction["_id"])
        
        return extraction
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching PDF extraction {extraction_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch extraction: {str(e)}")

@api_router.post("/pdf-processor/convert-to-project")
async def convert_extraction_to_project(
    extraction_id: str,
    project_metadata: dict,
    current_user: dict = Depends(get_current_user)
):
    """Convert PDF extraction data to a new project"""
    try:
        # Get extraction data
        extraction = await db.pdf_extractions.find_one({"id": extraction_id})
        if not extraction:
            raise HTTPException(status_code=404, detail="Extraction not found")
        
        extracted_data = extraction["extracted_data"]
        
        # Create BOQ items from line items
        boq_items = []
        for i, item in enumerate(extracted_data.get("line_items", [])):
            boq_item = {
                "serial_number": str(i + 1),
                "description": item.get("description", "Imported Item"),
                "unit": item.get("unit", "nos"),
                "quantity": item.get("quantity", 1.0),
                "rate": item.get("rate", 0.0),
                "amount": item.get("amount", item.get("quantity", 1.0) * item.get("rate", 0.0)),
                "gst_rate": 18.0,  # Default GST rate
                "gst_amount": (item.get("amount", item.get("quantity", 1.0) * item.get("rate", 0.0))) * 0.18,
                "total_with_gst": (item.get("amount", item.get("quantity", 1.0) * item.get("rate", 0.0))) * 1.18
            }
            boq_items.append(boq_item)
        
        # Calculate totals
        total_project_value = sum(item["total_with_gst"] for item in boq_items)
        
        # Create project
        project_data = {
            "id": str(uuid.uuid4()),
            "project_name": project_metadata.get("project_name", f"Project from PO {extracted_data.get('po_number', 'Unknown')}"),
            "architect": project_metadata.get("architect", extracted_data.get("vendor_name", "Unknown Architect")),
            "client_id": project_metadata.get("client_id", ""),
            "client_name": project_metadata.get("client_name", extracted_data.get("client_name", "Unknown Client")),
            "metadata": {
                "imported_from_pdf": True,
                "original_filename": extraction["original_filename"],
                "po_number": extracted_data.get("po_number"),
                "po_date": extracted_data.get("po_date"),
                "extraction_method": extracted_data.get("extraction_method"),
                "confidence_score": extracted_data.get("confidence_score"),
                **project_metadata.get("additional_metadata", {})
            },
            "boq_items": boq_items,
            "total_project_value": total_project_value,
            "advance_received": 0.0,
            "pending_payment": total_project_value,
            "created_by": current_user["id"],
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        # Create the project
        await db.projects.insert_one(project_data)
        
        # Update extraction record to mark as converted
        await db.pdf_extractions.update_one(
            {"id": extraction_id},
            {"$set": {
                "converted_to_project": True,
                "project_id": project_data["id"],
                "conversion_date": datetime.utcnow()
            }}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "pdf_converted_to_project", 
            f"Converted PDF extraction to project: {project_data['project_name']} (₹{total_project_value:,.2f})"
        )
        
        return {
            "message": "Project created successfully from PDF extraction",
            "project_id": project_data["id"],
            "project_name": project_data["project_name"],
            "total_value": total_project_value,
            "items_count": len(boq_items)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error converting extraction to project: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to convert to project: {str(e)}")

# Admin Configuration System
class WorkflowConfig(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    workflow_name: str
    workflow_type: str  # 'approval', 'billing', 'project', 'invoice'
    steps: List[Dict[str, Any]]
    roles_permissions: Dict[str, List[str]]
    notifications_config: Dict[str, Any]
    active: bool = True
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class SystemConfig(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    config_category: str  # 'ui', 'business', 'integration', 'notification'
    config_key: str
    config_value: Any
    config_type: str  # 'string', 'number', 'boolean', 'object', 'array'
    description: Optional[str] = None
    is_sensitive: bool = False
    requires_restart: bool = False
    created_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

@api_router.post("/admin/workflows", response_model=dict)
async def create_workflow_config(
    workflow_data: WorkflowConfig,
    current_user: dict = Depends(get_current_user)
):
    """Create a new workflow configuration"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can configure workflows")
        
        workflow_data.created_by = current_user["id"]
        workflow_data.updated_at = datetime.utcnow()
        
        await db.workflow_configs.insert_one(workflow_data.dict())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "workflow_created", 
            f"Created workflow configuration: {workflow_data.workflow_name} ({workflow_data.workflow_type})"
        )
        
        return {"message": "Workflow configuration created successfully", "workflow_id": workflow_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating workflow config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create workflow: {str(e)}")

@api_router.get("/admin/workflows")
async def get_workflow_configs(
    workflow_type: Optional[str] = None,
    active_only: bool = True,
    current_user: dict = Depends(get_current_user)
):
    """Get workflow configurations"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can view workflow configs")
        
        query = {}
        if workflow_type:
            query["workflow_type"] = workflow_type
        if active_only:
            query["active"] = True
        
        workflows = await db.workflow_configs.find(query).sort("created_at", -1).to_list(100)
        
        # Convert ObjectIds to strings for JSON serialization
        for workflow in workflows:
            if "_id" in workflow:
                workflow["_id"] = str(workflow["_id"])
        
        return workflows
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching workflow configs: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch workflows: {str(e)}")

@api_router.put("/admin/workflows/{workflow_id}")
async def update_workflow_config(
    workflow_id: str,
    update_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update workflow configuration"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can update workflows")
        
        workflow = await db.workflow_configs.find_one({"id": workflow_id})
        if not workflow:
            raise HTTPException(status_code=404, detail="Workflow configuration not found")
        
        update_data["updated_at"] = datetime.utcnow()
        
        await db.workflow_configs.update_one(
            {"id": workflow_id},
            {"$set": update_data}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "workflow_updated", 
            f"Updated workflow configuration: {workflow['workflow_name']}"
        )
        
        return {"message": "Workflow configuration updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating workflow config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update workflow: {str(e)}")

@api_router.post("/admin/system-config", response_model=dict)
async def create_system_config(
    config_data: SystemConfig,
    current_user: dict = Depends(get_current_user)
):
    """Create system configuration"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can configure system")
        
        config_data.created_by = current_user["id"]
        config_data.updated_at = datetime.utcnow()
        
        await db.system_configs.insert_one(config_data.dict())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "system_config_created", 
            f"Created system config: {config_data.config_category}.{config_data.config_key}"
        )
        
        return {"message": "System configuration created successfully", "config_id": config_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating system config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create system config: {str(e)}")

@api_router.get("/admin/system-config")
async def get_system_configs(
    category: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get system configurations"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can view system configs")
        
        query = {}
        if category:
            query["config_category"] = category
        
        configs = await db.system_configs.find(query).sort("config_category", 1).to_list(1000)
        
        # Group by category for better organization
        grouped_configs = {}
        for config in configs:
            # Convert ObjectIds to strings for JSON serialization
            if "_id" in config:
                config["_id"] = str(config["_id"])
                
            cat = config["config_category"]
            if cat not in grouped_configs:
                grouped_configs[cat] = []
            
            # Remove sensitive values
            if config.get("is_sensitive", False):
                config["config_value"] = "***HIDDEN***"
            
            grouped_configs[cat].append(config)
        
        return grouped_configs
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching system configs: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch system configs: {str(e)}")

@api_router.put("/admin/system-config/{config_id}")
async def update_system_config(
    config_id: str,
    update_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update system configuration"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can update system configs")
        
        config = await db.system_configs.find_one({"id": config_id})
        if not config:
            raise HTTPException(status_code=404, detail="System configuration not found")
        
        update_data["updated_at"] = datetime.utcnow()
        
        await db.system_configs.update_one(
            {"id": config_id},
            {"$set": update_data}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "system_config_updated", 
            f"Updated system config: {config['config_category']}.{config['config_key']}"
        )
        
        restart_required = config.get("requires_restart", False) or update_data.get("requires_restart", False)
        
        return {
            "message": "System configuration updated successfully",
            "restart_required": restart_required
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating system config: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update system config: {str(e)}")

# Company Profile Management Endpoints
@api_router.post("/company-profiles", response_model=dict)
async def create_company_profile(
    profile_data: CompanyProfile,
    current_user: dict = Depends(get_current_user)
):
    """Create a new company profile - Admin/Super Admin only"""
    try:
        if current_user["role"] not in [UserRole.ADMIN, UserRole.SUPER_ADMIN]:
            raise HTTPException(status_code=403, detail="Only admin/super admin can create company profiles")
        
        profile_data.created_by = current_user["id"]
        profile_data.updated_at = datetime.utcnow()
        
        # Ensure only one default location and bank
        default_location_count = sum(1 for loc in profile_data.locations if loc.is_default)
        default_bank_count = sum(1 for bank in profile_data.bank_details if bank.is_default)
        
        if default_location_count > 1:
            raise HTTPException(status_code=400, detail="Only one location can be set as default")
        if default_bank_count > 1:
            raise HTTPException(status_code=400, detail="Only one bank account can be set as default")
        
        # Set default IDs if defaults are specified
        for loc in profile_data.locations:
            if loc.is_default:
                profile_data.default_location_id = loc.id
        for bank in profile_data.bank_details:
            if bank.is_default:
                profile_data.default_bank_id = bank.id
        
        await db.company_profiles.insert_one(profile_data.dict())
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "company_profile_created", 
            f"Created company profile: {profile_data.company_name} with {len(profile_data.locations)} locations and {len(profile_data.bank_details)} bank accounts"
        )
        
        return {"message": "Company profile created successfully", "profile_id": profile_data.id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating company profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create company profile: {str(e)}")

@api_router.get("/company-profiles")
async def get_company_profiles(
    current_user: dict = Depends(get_current_user)
):
    """Get all company profiles"""
    try:
        profiles = await db.company_profiles.find().sort("created_at", -1).to_list(100)
        return profiles
        
    except Exception as e:
        logger.error(f"Error fetching company profiles: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch company profiles: {str(e)}")

@api_router.get("/company-profiles/{profile_id}")
async def get_company_profile(
    profile_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get specific company profile"""
    try:
        profile = await db.company_profiles.find_one({"id": profile_id})
        if not profile:
            raise HTTPException(status_code=404, detail="Company profile not found")
        
        return profile
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching company profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch company profile: {str(e)}")

@api_router.put("/company-profiles/{profile_id}")
async def update_company_profile(
    profile_id: str,
    update_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update company profile - Admin/Super Admin only"""
    try:
        if current_user["role"] not in [UserRole.ADMIN, UserRole.SUPER_ADMIN]:
            raise HTTPException(status_code=403, detail="Only admin/super admin can update company profiles")
        
        profile = await db.company_profiles.find_one({"id": profile_id})
        if not profile:
            raise HTTPException(status_code=404, detail="Company profile not found")
        
        update_data["updated_at"] = datetime.utcnow()
        
        # Validate defaults if provided
        if "locations" in update_data:
            default_location_count = sum(1 for loc in update_data["locations"] if loc.get("is_default", False))
            if default_location_count > 1:
                raise HTTPException(status_code=400, detail="Only one location can be set as default")
        
        if "bank_details" in update_data:
            default_bank_count = sum(1 for bank in update_data["bank_details"] if bank.get("is_default", False))
            if default_bank_count > 1:
                raise HTTPException(status_code=400, detail="Only one bank account can be set as default")
        
        await db.company_profiles.update_one(
            {"id": profile_id},
            {"$set": update_data}
        )
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "company_profile_updated", 
            f"Updated company profile: {profile.get('company_name', 'Unknown')}"
        )
        
        return {"message": "Company profile updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating company profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update company profile: {str(e)}")

@api_router.delete("/company-profiles/{profile_id}")
async def delete_company_profile(
    profile_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete company profile - Super Admin only"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can delete company profiles")
        
        profile = await db.company_profiles.find_one({"id": profile_id})
        if not profile:
            raise HTTPException(status_code=404, detail="Company profile not found")
        
        # Check if profile is being used in any projects
        projects_using_profile = await db.projects.count_documents({"company_profile_id": profile_id})
        if projects_using_profile > 0:
            raise HTTPException(status_code=400, detail=f"Cannot delete company profile. It is being used in {projects_using_profile} projects")
        
        await db.company_profiles.delete_one({"id": profile_id})
        
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "company_profile_deleted", 
            f"Deleted company profile: {profile.get('company_name', 'Unknown')}"
        )
        
        return {"message": "Company profile deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting company profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete company profile: {str(e)}")

# Project Metadata Validation Engine
class ProjectMetadataValidator:
    """Validates project metadata against BOQ data"""
    
    def __init__(self):
        self.validation_errors = []
    
    async def validate_project_metadata(self, metadata: List[dict], boq_items: List[dict]) -> dict:
        """Validate project metadata against BOQ items"""
        self.validation_errors = []
        
        if not metadata or len(metadata) == 0:
            self.validation_errors.append("At least one Purchase Order entry is required")
            return {"valid": False, "errors": self.validation_errors}
        
        # Check mandatory fields
        for i, po_entry in enumerate(metadata):
            if not po_entry.get("purchase_order_number"):
                self.validation_errors.append(f"Row {i+1}: Purchase Order Number is mandatory")
        
        # Calculate totals from BOQ
        boq_total = sum(item.get("total_with_gst", 0) for item in boq_items)
        
        # Calculate totals from metadata
        metadata_total = sum(po_entry.get("po_inv_value", 0) for po_entry in metadata)
        
        # Validate total value match (allow 5% variance)
        if abs(boq_total - metadata_total) > (boq_total * 0.05):
            self.validation_errors.append(
                f"Total value mismatch: BOQ total (₹{boq_total:,.2f}) vs PO total (₹{metadata_total:,.2f}). "
                f"Difference: ₹{abs(boq_total - metadata_total):,.2f}"
            )
        
        # Validate percentage calculations
        for i, po_entry in enumerate(metadata):
            po_value = po_entry.get("po_inv_value", 0)
            if po_value > 0:
                # Validate ABG percentage
                abg_percent = po_entry.get("abg_percentage", 0)
                if abg_percent < 0 or abg_percent > 100:
                    self.validation_errors.append(f"Row {i+1}: ABG percentage must be between 0-100%")
                
                # Validate RA Bill percentage
                ra_percent = po_entry.get("ra_bill_with_taxes_percentage", 0)
                if ra_percent < 0 or ra_percent > 100:
                    self.validation_errors.append(f"Row {i+1}: RA Bill percentage must be between 0-100%")
                
                # Validate Erection percentage
                erection_percent = po_entry.get("erection_percentage", 0)
                if erection_percent < 0 or erection_percent > 100:
                    self.validation_errors.append(f"Row {i+1}: Erection percentage must be between 0-100%")
                
                # Validate PBG percentage
                pbg_percent = po_entry.get("pbg_percentage", 0)
                if pbg_percent < 0 or pbg_percent > 100:
                    self.validation_errors.append(f"Row {i+1}: PBG percentage must be between 0-100%")
        
        return {
            "valid": len(self.validation_errors) == 0,
            "errors": self.validation_errors,
            "boq_total": boq_total,
            "metadata_total": metadata_total,
            "variance": abs(boq_total - metadata_total)
        }

@api_router.post("/projects/validate-metadata")
async def validate_project_metadata(
    validation_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Validate project metadata against BOQ data"""
    try:
        metadata = validation_data.get("metadata", [])
        boq_items = validation_data.get("boq_items", [])
        
        validator = ProjectMetadataValidator()
        result = await validator.validate_project_metadata(metadata, boq_items)
        
        return result
        
    except Exception as e:
        logger.error(f"Error validating project metadata: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to validate project metadata: {str(e)}")

@api_router.post("/admin/clear-database")
async def clear_database(
    confirmation: dict,
    current_user: dict = Depends(get_current_user)
):
    """Clear entire database - DANGEROUS OPERATION - Super Admin Only"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can clear database")
        
        # Require explicit confirmation
        if not confirmation.get("confirm_clear") or confirmation.get("confirmation_text") != "DELETE ALL DATA":
            raise HTTPException(status_code=400, detail="Proper confirmation required to clear database")
        
        # Get collection stats before clearing
        collections_to_clear = [
            "projects", "invoices", "clients", "bank_guarantees", 
            "pdf_extractions", "master_items", "workflow_configs", 
            "system_configs", "activity_logs"
        ]
        
        stats_before = {}
        for collection_name in collections_to_clear:
            try:
                count = await db[collection_name].count_documents({})
                stats_before[collection_name] = count
            except:
                stats_before[collection_name] = 0
        
        # Clear all collections except users (preserve user accounts)
        cleared_collections = []
        total_deleted = 0
        
        for collection_name in collections_to_clear:
            try:
                result = await db[collection_name].delete_many({})
                cleared_collections.append({
                    "collection": collection_name,
                    "deleted_count": result.deleted_count,
                    "previous_count": stats_before.get(collection_name, 0)
                })
                total_deleted += result.deleted_count
            except Exception as e:
                logger.error(f"Error clearing collection {collection_name}: {str(e)}")
                cleared_collections.append({
                    "collection": collection_name,
                    "deleted_count": 0,
                    "previous_count": stats_before.get(collection_name, 0),
                    "error": str(e)
                })
        
        # Log this critical action
        await log_activity(
            current_user["id"], current_user["email"], current_user["role"],
            "database_cleared", 
            f"🚨 CRITICAL: Database cleared by super admin. Total records deleted: {total_deleted}. Collections cleared: {len(cleared_collections)}"
        )
        
        return {
            "message": "Database cleared successfully",
            "timestamp": datetime.utcnow(),
            "cleared_by": {
                "user_id": current_user["id"],
                "email": current_user["email"]
            },
            "statistics": {
                "total_records_deleted": total_deleted,
                "collections_cleared": len([c for c in cleared_collections if c.get("deleted_count", 0) > 0]),
                "collections_details": cleared_collections
            },
            "preserved": {
                "users": "User accounts preserved for system access"
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error clearing database: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clear database: {str(e)}")

@api_router.get("/admin/system-health")
async def get_system_health(current_user: dict = Depends(get_current_user)):
    """Get system health and status information"""
    try:
        if current_user["role"] != UserRole.SUPER_ADMIN:
            raise HTTPException(status_code=403, detail="Only super admin can view system health")
        
        # Database connectivity check
        db_status = "healthy"
        try:
            await db.command("ping")
        except:
            db_status = "unhealthy"
        
        # Collection counts
        collections_status = {}
        collection_names = ["users", "projects", "invoices", "clients", "activity_logs", "master_items", "bank_guarantees", "pdf_extractions"]
        
        for collection_name in collection_names:
            try:
                count = await db[collection_name].count_documents({})
                collections_status[collection_name] = {"count": count, "status": "healthy"}
            except Exception as e:
                collections_status[collection_name] = {"count": 0, "status": "error", "error": str(e)}
        
        # Recent activity
        recent_logs = await db.activity_logs.find().sort("timestamp", -1).limit(10).to_list(10)
        
        # Convert ObjectId to string for JSON serialization
        for log in recent_logs:
            if "_id" in log:
                log["_id"] = str(log["_id"])
        
        # System info
        health_info = {
            "database": {
                "status": db_status,
                "collections": collections_status
            },
            "application": {
                "version": "1.0.0",
                "uptime": "Available",  # Could be calculated from startup time
                "environment": os.environ.get("ENVIRONMENT", "development")
            },
            "recent_activity": recent_logs,
            "timestamp": datetime.utcnow()
        }
        
        return health_info
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting system health: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get system health: {str(e)}")

# Include router
app.include_router(api_router)

@app.on_event("startup")
async def startup_event():
    await init_super_admin()
    logger.info("Application started successfully")

@app.on_event("shutdown")
async def shutdown_event():
    client.close()
    logger.info("Application shutdown")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)